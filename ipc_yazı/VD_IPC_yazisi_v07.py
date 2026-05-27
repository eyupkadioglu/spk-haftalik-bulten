# -*- coding: utf-8 -*-
"""
Vergi Dairesi Yazısı Word oluşturucu
Copyright: Eyüp Kadıoğlu 2026

Kullanım:
1) Bu dosyayı SPKYTMIPC.xlsx dosyasının bulunduğu klasöre koyun.
2) Aynı klasörde Word şablonları da bulunsun: sablonlar/ klasöründe
   - vergidaireipctahsilatsorma_vd.docx
   - vergidaireipctahsilatsorma_spk.docx
3) Gerekirse şu paketleri kurun:
   pip install python-docx openpyxl
4) Çalıştırın:
   python VD_IPC_yazisi_v07.py

Not:
- Vergi Dairesi yazısı için SPKYTMIPC.xlsx dosyasının "Kurtarılan_Sayfa1" sayfasından veri alınır.
- Word'deki sarı alanlar üretilen dosyada varsayılan olarak sarı bırakılmaz.
"""

from __future__ import annotations

import os
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.shared import Pt
except ImportError:
    raise SystemExit("Eksik paket: python-docx. Kurulum: pip install python-docx")

try:
    from openpyxl import Workbook, load_workbook
except ImportError:
    raise SystemExit("Eksik paket: openpyxl. Kurulum: pip install openpyxl")

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
except ImportError:
    raise SystemExit("tkinter bulunamadı. Windows Python kurulumunda tkinter genelde hazır gelir.")


# =========================
# AYARLAR
# =========================

EXCEL_DOSYA_ADI = "SPKYTMIPC.xlsx"
KALAN_ALACAK_DOSYA_ADI = "Kalan Alacak Dosyaları Dökümü.xlsx"
SABLON_KLASOR_ADI = "sablonlar"
VERGI_DAIRESI_VD_SABLON_ADI = "vergidaireipctahsilatsorma_vd.docx"
VERGI_DAIRESI_SPK_SABLON_ADI = "vergidaireipctahsilatsorma_spk.docx"
VERGI_DAIRESI_SHEET_ADI = "Kurtarılan_Sayfa1"
CIKTI_KLASOR_ADI = "olusan_word_yazilari"
LOG_DOSYA_ADI = "log_kaydi.xlsx"
TUTARSIZLIK_YORUMU = "İPC Tutarı (TL) ve Ödenen Tutar (TL) bilgilerinde tutarsızlık var, kontrol ediniz."
ORAN_FARKLI_YORUMU = "Ödenen Tutar (TL)/İPC Tutarı (TL) oranı %75 veya %100 den farklıdır"
SABLON_KISI_CEZA_IFADELERI = (
    "ilgili gerçek kişi hakkında idari para cezası (İPC) verilmiş olup, SPK’dan elde edilen bilgiler çerçevesinde söz konusu idari para cezasının Müdürlüğünüze ödendiği tespit edilmiştir.",
    "ilgili şirket hakkında idari para cezası (İPC) verilmiş olup, SPK’dan elde edilen bilgiler çerçevesinde söz konusu idari para cezasının Müdürlüğünüze ödendiği tespit edilmiştir.",
)
SABLON_SON_PARAGRAF_TEKIL_CEZA_IFADESI = "Müdürlüğünüzce tahsil edilen idari para cezasının"
SABLON_SON_PARAGRAF_COGUL_CEZA_IFADESI = "Müdürlüğünüzce tahsil edilen idari para cezalarının"

# True: Word'deki sarı işaretler çıktı dosyasında temizlenir.
# False: Dinamik alanlar sarı vurgulu kalır.
FINALDE_SARI_ALANLARI_TEMIZLE = True
COPYRIGHT = "Eyüp Kadıoğlu 2026 Version 0.7"


@dataclass
class BelgeAyari:
    secim_adi: str
    excel_sayfa_adi: Optional[str]
    aktif: bool
    aciklama: str = ""


BELGE_TURU_AYARLARI: Dict[str, BelgeAyari] = {
    "VERGI_DAIRESI": BelgeAyari(
        secim_adi="Vergi Dairesi Yazısı",
        excel_sayfa_adi=VERGI_DAIRESI_SHEET_ADI,
        aktif=True,
        aciklama="Vergi Dairesi için İPC tahsilat yazısı",
    ),
}


# =========================
# VERGİ DAİRESİ FONKSİYONLARI
# =========================

def vergi_dairesi_id_ayikla(metin: str) -> str:
    metin = str(metin or "").strip()
    eslesme = re.search(r"\((\d+)\)", metin)
    if eslesme and len(eslesme.group(1)) in (10, 11):
        return eslesme.group(1)
    return ""


def vergi_dairesi_muhatap_ayikla(metin: str) -> str:
    metin = str(metin or "").strip()
    metin = re.sub(r"\s*-\s*\(\d+\)", "", metin).strip()
    metin = re.sub(r"\s*\(\d+\)", "", metin).strip()
    return metin


def girdi_bilgisi_ayikla(girdi: str) -> Dict[str, str]:
    metin = str(girdi or "").strip()
    kalan = metin
    bilgi = {"arama": metin, "ipc_id": "", "daire_kodu": "", "vd_kodu": ""}

    def etiketli_al(pattern: str, alan: str, numeric_daire: bool = False) -> None:
        nonlocal kalan
        eslesme = re.search(pattern, kalan, flags=re.IGNORECASE)
        if not eslesme:
            return
        deger = eslesme.group(1).strip()
        if numeric_daire and deger.isdigit():
            bilgi["vd_kodu"] = deger
        else:
            bilgi[alan] = deger
        kalan = (kalan[:eslesme.start()] + " " + kalan[eslesme.end():]).strip()

    etiketli_al(r"\b(?:IPC|İPC)\s*(?:ID|İD|NO|DOSYA\s*NO)?\s*[:=]?\s*([A-Za-z0-9./_-]+)", "ipc_id")
    etiketli_al(r"\b(?:DAIRE|DAİRE)\s*(?:KODU)?\s*[:=]?\s*([A-Za-zÇĞİÖŞÜçğıöşü0-9]+)", "daire_kodu", numeric_daire=True)
    etiketli_al(r"\bVD\s*(?:KODU)?\s*[:=]?\s*(\d{3,10})", "vd_kodu")

    tokenlar = re.findall(r"\S+", kalan)
    silinecek_indeksler = set()
    if not bilgi["ipc_id"]:
        for i, token in enumerate(tokenlar):
            temiz = token.strip(" ,;")
            if re.fullmatch(r"\d{1,9}", temiz):
                bilgi["ipc_id"] = temiz
                silinecek_indeksler.add(i)
                break

    if bilgi["ipc_id"] and not bilgi["daire_kodu"] and not bilgi["vd_kodu"]:
        for i, token in enumerate(tokenlar):
            if i in silinecek_indeksler:
                continue
            temiz = token.strip(" ,;")
            if re.fullmatch(r"\d{3,10}", temiz):
                bilgi["vd_kodu"] = temiz
                silinecek_indeksler.add(i)
                break
            if re.fullmatch(r"[A-Za-zÇĞİÖŞÜçğıöşü]{2,10}", temiz) and temiz == temiz.upper():
                bilgi["daire_kodu"] = temiz
                silinecek_indeksler.add(i)
                break

    if silinecek_indeksler:
        kalan = " ".join(token for i, token in enumerate(tokenlar) if i not in silinecek_indeksler)

    bilgi["arama"] = vergi_dairesi_muhatap_ayikla(kalan).strip() or metin
    return bilgi


def hucre_kodu_metni(value: Any) -> str:
    metin = excel_degeri_metni(value)
    if metin.endswith(".0"):
        metin = metin[:-2]
    return metin.strip()


def satir_filtreye_uygun_mu(ws, row: int, filtre: Dict[str, str]) -> bool:
    ipc_id = str(filtre.get("ipc_id", "")).strip()
    daire_kodu = arama_metni_normalize(filtre.get("daire_kodu", ""))
    vd_kodu = re.sub(r"\D", "", str(filtre.get("vd_kodu", "")))

    if ipc_id:
        satir_ipc_dosya_no = hucre_kodu_metni(ws.cell(row=row, column=1).value)
        satir_detay_id = hucre_kodu_metni(ws.cell(row=row, column=10).value)
        if ipc_id.upper() not in {satir_ipc_dosya_no.upper(), satir_detay_id.upper()}:
            return False

    if daire_kodu:
        satir_daire = arama_metni_normalize(hucre_metni(ws.cell(row=row, column=2)))
        if satir_daire != daire_kodu:
            return False

    if vd_kodu:
        satir_vd_kodu = re.sub(r"\D", "", hucre_kodu_metni(ws.cell(row=row, column=12).value))
        if satir_vd_kodu != vd_kodu:
            return False

    return True


def muhatap_gorunum_formatla(muhatap: str, kimlik_no: str) -> str:
    if kimlik_turu_belirle(kimlik_no) == "TCKN":
        return ad_soyad_formatla(muhatap)
    return unvan_bas_harf_formatla(muhatap)


def unvan_bas_harf_formatla(metin: str) -> str:
    parcalar = re.split(r"(\s+)", str(metin or "").strip())
    return "".join(turkce_bas_harf_buyut(parca) if parca.strip() else parca for parca in parcalar)


def konu_muhatap_kisa_formatla(muhatap: str, kimlik_no: str) -> str:
    parcalar = re.split(r"\s+", str(muhatap or "").strip())
    if not parcalar:
        return ""

    if kimlik_turu_belirle(kimlik_no) == "TCKN":
        if len(parcalar) == 1:
            return turkce_buyuk_harf(parcalar[0])
        ad_ilk_harf = turkce_buyuk_harf(parcalar[0][0])
        soyad = turkce_buyuk_harf(parcalar[-1])
        return f"{ad_ilk_harf}. {soyad}"

    return " ".join(turkce_bas_harf_buyut(parca) for parca in parcalar[:2])


def konu_metni_olustur(veri: Dict[str, Any]) -> str:
    muhataplar = []
    gorulen = set()
    for kayit in [veri] + list(veri.get("birlesik_kayitlar", [])):
        kisa = konu_muhatap_kisa_formatla(kayit.get("muhatap", ""), kayit.get("kimlik_no", ""))
        anahtar = arama_metni_normalize(kisa)
        if kisa and anahtar not in gorulen:
            muhataplar.append(kisa)
            gorulen.add(anahtar)

    parantez = ", ".join(muhataplar)
    return f"İdari Para Cazası Hk ({parantez})" if parantez else "İdari Para Cazası Hk"


def kisi_turu_sayilari(veri: Dict[str, Any]) -> Tuple[int, int]:
    gercek_kisiler = set()
    tuzel_kisiler = set()
    for kayit in [veri] + list(veri.get("birlesik_kayitlar", [])):
        kimlik_no = str(kayit.get("kimlik_no", "")).strip()
        anahtar = kimlik_no or arama_metni_normalize(kayit.get("muhatap", ""))
        if kimlik_turu_belirle(kimlik_no) == "TCKN":
            gercek_kisiler.add(anahtar)
        else:
            tuzel_kisiler.add(anahtar)
    return len(gercek_kisiler), len(tuzel_kisiler)


def tablo_muhatap_turu_sayilari(veri: Dict[str, Any]) -> Tuple[int, int]:
    gercek_kisiler = set()
    tuzel_kisiler = set()
    for kayit in [veri] + list(veri.get("birlesik_kayitlar", [])):
        kimlik_no = str(kayit.get("kimlik_no", "")).strip()
        anahtar = kimlik_no or arama_metni_normalize(kayit.get("muhatap", ""))
        if not anahtar:
            continue
        if kimlik_turu_belirle(kimlik_no) == "TCKN":
            gercek_kisiler.add(anahtar)
        else:
            tuzel_kisiler.add(anahtar)
    return len(gercek_kisiler), len(tuzel_kisiler)


def kisi_ceza_ifadesi_olustur(veri: Dict[str, Any]) -> str:
    kisi_ifadesi, ceza_ifadesi, ceza_tamlama = kisi_ceza_ifadesi_parcalari(veri)
    return (
        f"ilgili {kisi_ifadesi} hakkında {ceza_ifadesi} (İPC) verilmiş olup, "
        f"SPK’dan elde edilen bilgiler çerçevesinde söz konusu {ceza_tamlama} "
        "Müdürlüğünüze ödendiği tespit edilmiştir."
    )


def kisi_ceza_ifadesi_parcalari(veri: Dict[str, Any]) -> Tuple[str, str, str]:
    gercek_sayisi, tuzel_sayisi = tablo_muhatap_turu_sayilari(veri)
    toplam = gercek_sayisi + tuzel_sayisi

    if gercek_sayisi and tuzel_sayisi:
        kisi_ifadesi = "gerçek ve tüzel kişiler"
    elif gercek_sayisi:
        kisi_ifadesi = "gerçek kişi" if gercek_sayisi == 1 else "gerçek kişiler"
    else:
        kisi_ifadesi = "tüzel kişi" if tuzel_sayisi <= 1 else "tüzel kişiler"

    ceza_ifadesi = "idari para cezası" if toplam <= 1 else "idari para cezaları"
    ceza_tamlama = "idari para cezasının" if toplam <= 1 else "idari para cezalarının"
    return kisi_ifadesi, ceza_ifadesi, ceza_tamlama


def kisi_ceza_ifadesini_guncelle(doc: Document, veri: Dict[str, Any]) -> None:
    yeni_ifade = kisi_ceza_ifadesi_olustur(veri)
    kisi_ifadesi, ceza_ifadesi, ceza_tamlama = kisi_ceza_ifadesi_parcalari(veri)
    for p in belgedeki_paragraflar(doc):
        for eski_ifade in SABLON_KISI_CEZA_IFADELERI:
            if eski_ifade in p.text:
                yeni_metin = p.text.replace(eski_ifade, yeni_ifade)
                paragraf_yaz(p, yeni_metin, fmt=paragraf_ilk_run_format(p), sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
                return
        if "hakkında idari para cezası" in p.text and "tespit edilmiştir" in p.text:
            yeni_metin = p.text
            yeni_metin = yeni_metin.replace(
                "gerçek kişi hakkında idari para cezası",
                f"{kisi_ifadesi} hakkında {ceza_ifadesi}",
            )
            yeni_metin = yeni_metin.replace(
                "şirket hakkında idari para cezası",
                f"{kisi_ifadesi} hakkında {ceza_ifadesi}",
            )
            yeni_metin = yeni_metin.replace("söz konusu idari para cezasının", f"söz konusu {ceza_tamlama}")
            if yeni_metin != p.text:
                paragraf_yaz(p, yeni_metin, fmt=paragraf_ilk_run_format(p), sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
                return


def son_paragraf_ceza_ifadesini_guncelle(doc: Document, veri: Dict[str, Any]) -> None:
    yeni_ifade = (
        SABLON_SON_PARAGRAF_COGUL_CEZA_IFADESI
        if len(belge_tablo_kayitlari(veri)) > 1
        else SABLON_SON_PARAGRAF_TEKIL_CEZA_IFADESI
    )
    for p in belgedeki_paragraflar(doc):
        yeni_metin = p.text
        yeni_metin = yeni_metin.replace(SABLON_SON_PARAGRAF_TEKIL_CEZA_IFADESI, yeni_ifade)
        yeni_metin = yeni_metin.replace(SABLON_SON_PARAGRAF_COGUL_CEZA_IFADESI, yeni_ifade)
        if yeni_metin != p.text:
            paragraf_yaz(p, yeni_metin, fmt=paragraf_ilk_run_format(p), sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)


def vergi_dairesi_hitap_formatla(vergi_dairesi_adi: str) -> str:
    ad = turkce_buyuk_harf(str(vergi_dairesi_adi or "").strip())
    ad = re.sub(r"\s*-\s*\(\d+\)", " ", ad)
    ad = re.sub(r"\s*\(\d+\)", " ", ad)
    ad = re.sub(r"\s+", " ", ad).strip(" -")
    if re.search(r"\bMALM[ÜU]D[ÜU]RL[ÜU][ĞG][ÜU]\b", ad):
        return ad
    if re.search(r"\bDEFTERDARLI[ĞG]I\b", ad):
        return ad
    ad = re.sub(r"^\s*\d+\s*[-/]\s*", " ", ad)
    ad = re.sub(r"\s*[-/]\s*\d+\s*$", " ", ad)
    ad = re.sub(r"\bV\.?\s*D\.?\b", " ", ad)
    ad = re.sub(r"\bVERG[İI]\s+DA[İI]RES[İI](?:\s+M[ÜU]D[ÜU]RL[ÜU]Ğ[ÜU])?\b", " ", ad)
    ad = re.sub(r"\s+", " ", ad).strip(" -")
    if not ad:
        return "VERGİ DAİRESİ MÜDÜRLÜĞÜ"
    return f"{ad} VERGİ DAİRESİ MÜDÜRLÜĞÜ"


def belgedeki_paragraflar(doc) -> List:
    """Belgedeki tüm paragrafları döndürür."""
    sonuc = []
    for p in doc.paragraphs:
        sonuc.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    sonuc.append(p)
    return sonuc


def hucre_yaz(cell, text: str, alignment=None, font_size=None) -> None:
    """Tablo hücresini temizleyip yeni metin yazar."""
    cell.text = ""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    if font_size is not None:
        run.font.size = font_size


def tablo_satiri_yaz(row, veri: Dict[str, str]) -> None:
    if len(row.cells) < 6:
        raise ValueError("Vergi Dairesi tablosunda yeterli hücre yok.")
    hucre_yaz(row.cells[0], muhatap_gorunum_formatla(veri["muhatap"], veri["kimlik_no"]))
    hucre_yaz(row.cells[1], veri["kimlik_no"])
    hucre_yaz(row.cells[2], f"{veri['karar_tarihi']}-{veri['karar_no']}")
    hucre_yaz(row.cells[3], veri["pc_tutari"], WD_ALIGN_PARAGRAPH.RIGHT)
    hucre_yaz(row.cells[4], odenen_tutari_word_icin_formatla(veri["odenen_tutar"]), WD_ALIGN_PARAGRAPH.RIGHT)
    hucre_yaz(row.cells[5], veri["tahsilat_tarihi_fis_no"], WD_ALIGN_PARAGRAPH.CENTER)


def ipc_odenen_tutar_tutarsiz_mi(veri: Dict[str, Any]) -> bool:
    ipc_tutari = para_decimal(veri.get("pc_tutari", ""))
    odenen_tutar = para_decimal(veri.get("odenen_tutar", ""))
    return ipc_tutari is not None and odenen_tutar is not None and ipc_tutari < odenen_tutar


def tablo_tutarsiz_satir_indeksleri(tablo_kayitlari: List[Dict[str, Any]]) -> List[int]:
    return [
        grup["satirlar"][0]
        for grup in tablo_odeme_gruplari(tablo_kayitlari).values()
        if grup["ipc_tutari"] < grup["odenen_toplam"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    ]


def oran_farkli_satir_indeksleri(tablo_kayitlari: List[Dict[str, Any]]) -> List[int]:
    sonuc = []
    for grup in tablo_odeme_gruplari(tablo_kayitlari).values():
        ipc_tutari = grup["ipc_tutari"]
        if ipc_tutari == Decimal("0.00"):
            continue
        odenen_toplam = grup["odenen_toplam"].quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        oran = (odenen_toplam / ipc_tutari).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        if oran not in (Decimal("0.75"), Decimal("1.00")):
            sonuc.append(grup["satirlar"][0])
    return sonuc


def tablo_tutarsizlik_yorumlarini_ekle(doc: Document, table, tablo_kayitlari: List[Dict[str, Any]]) -> None:
    yorum_eklenenler = set()
    yorumlar = [
        (row_idx, TUTARSIZLIK_YORUMU)
        for row_idx in tablo_tutarsiz_satir_indeksleri(tablo_kayitlari)
    ]
    yorumlar.extend(
        (row_idx, ORAN_FARKLI_YORUMU)
        for row_idx in oran_farkli_satir_indeksleri(tablo_kayitlari)
    )

    for row_idx, yorum_metni in yorumlar:
        if row_idx >= len(table.rows):
            continue

        cell = table.rows[row_idx].cells[4]
        yorum_anahtari = (id(cell._tc), yorum_metni)
        if yorum_anahtari in yorum_eklenenler:
            continue

        runs = [run for p in cell.paragraphs for run in p.runs if run.text]
        if not runs:
            runs = [cell.paragraphs[0].add_run(cell.text or " ")]

        doc.add_comment(runs, yorum_metni, author="Kontrol", initials="")
        yorum_eklenenler.add(yorum_anahtari)


def tablo_kayit_grup_anahtari(tablo_veri: Dict[str, Any]) -> Tuple[str, str, str, str, str]:
    ipc_tutari = para_decimal(tablo_veri.get("pc_tutari", ""))
    return (
        str(tablo_veri.get("muhatap", "")),
        str(tablo_veri.get("kimlik_no", "")),
        str(tablo_veri.get("karar_tarihi", "")),
        str(tablo_veri.get("karar_no", "")),
        decimal_para_metni(ipc_tutari) if ipc_tutari is not None else str(tablo_veri.get("pc_tutari", "")),
    )


def tablo_odeme_gruplari(tablo_kayitlari: List[Dict[str, Any]]) -> Dict[Tuple[str, str, str, str, str], Dict[str, Any]]:
    gruplar: Dict[Tuple[str, str, str, str, str], Dict[str, Any]] = {}
    for row_idx, tablo_veri in enumerate(tablo_kayitlari, start=1):
        ipc_tutari = para_decimal(tablo_veri.get("pc_tutari", ""))
        odenen_tutar = para_decimal(tablo_veri.get("odenen_tutar", ""))
        if ipc_tutari is None or odenen_tutar is None:
            continue

        anahtar = tablo_kayit_grup_anahtari(tablo_veri)
        grup = gruplar.setdefault(
            anahtar,
            {"ipc_tutari": ipc_tutari, "odenen_toplam": Decimal("0.00"), "satirlar": []},
        )
        grup["odenen_toplam"] += odenen_tutar
        grup["satirlar"].append(row_idx)
    return gruplar


def belge_tablo_kayitlari(veri: Dict[str, Any]) -> List[Dict[str, Any]]:
    kayitlar = [veri]
    kayitlar.extend(veri.get("parcali_tahsilatlar", []))
    kayitlar.extend(veri.get("ek_ipc_kayitlari", []))
    for ek_veri in veri.get("birlesik_kayitlar", []):
        kayitlar.append(ek_veri)
        kayitlar.extend(ek_veri.get("parcali_tahsilatlar", []))
        kayitlar.extend(ek_veri.get("ek_ipc_kayitlari", []))
    return kayitlar


def ayni_hucreleri_dikey_birlestir(table, baslangic_satiri: int = 1) -> None:
    if len(table.rows) <= baslangic_satiri + 1:
        return

    sutun_sayisi = len(table.rows[baslangic_satiri].cells)
    for col in range(sutun_sayisi):
        if col in (4, 5):
            continue

        grup_baslangic = baslangic_satiri
        onceki_metin = table.rows[baslangic_satiri].cells[col].text.strip()

        for row_idx in range(baslangic_satiri + 1, len(table.rows) + 1):
            metin = ""
            if row_idx < len(table.rows):
                metin = table.rows[row_idx].cells[col].text.strip()

            if row_idx < len(table.rows) and metin == onceki_metin and metin:
                continue

            if row_idx - grup_baslangic > 1 and onceki_metin:
                birlesik_hucre = table.rows[grup_baslangic].cells[col].merge(table.rows[row_idx - 1].cells[col])
                alignment = WD_ALIGN_PARAGRAPH.RIGHT if col in (3, 4) else None
                hucre_yaz(birlesik_hucre, onceki_metin, alignment)

            grup_baslangic = row_idx
            onceki_metin = metin


def tum_sari_vurgulari_temizle(doc) -> None:
    """Belgedeki tüm sarı vurguları temizler."""
    for p in belgedeki_paragraflar(doc):
        for run in p.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                run.font.highlight_color = None


def sari_paragraflari_bul(doc: Document) -> List:
    paragraflar = []

    for p in doc.paragraphs:
        if paragrafta_sari_var_mi(p):
            paragraflar.append(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if paragrafta_sari_var_mi(p):
                        paragraflar.append(p)

    return paragraflar


def paragrafta_sari_var_mi(paragraph) -> bool:
    return any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in paragraph.runs)


def paragraf_ilk_run_format(paragraph) -> Dict[str, object]:
    for run in paragraph.runs:
        if run.text:
            return run_format_snapshot(run)
    if paragraph.runs:
        return run_format_snapshot(paragraph.runs[0])
    # Güvenli varsayılan
    temp = paragraph.add_run("")
    fmt = run_format_snapshot(temp)
    paragraph._p.remove(temp._r)
    return fmt


def run_format_snapshot(run) -> Dict[str, object]:
    """Run'un format bilgilerini snapshot alır."""
    return {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "style": run.style,
        "font_name": run.font.name,
        "font_size": run.font.size,
        "highlight": run.font.highlight_color,
        "color_rgb": None,
    }


def format_uygula(run, fmt: Dict[str, object], sariyi_temizle: bool = True) -> None:
    run.bold = fmt.get("bold")
    run.italic = fmt.get("italic")
    run.underline = fmt.get("underline")
    try:
        run.style = fmt.get("style")
    except Exception:
        pass
    run.font.name = fmt.get("font_name")
    run.font.size = fmt.get("font_size")
    if fmt.get("color_rgb") is not None:
        run.font.color.rgb = fmt.get("color_rgb")
    if sariyi_temizle:
        run.font.highlight_color = None
    else:
        run.font.highlight_color = fmt.get("highlight")


def paragrafi_temizle(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def paragraf_yaz(paragraph, text: str, fmt: Optional[Dict[str, object]] = None, sariyi_temizle: bool = True) -> None:
    """Paragrafın metnini değiştirir, paragraf stilini/yerleşimini korur."""
    if fmt is None:
        fmt = paragraf_ilk_run_format(paragraph)

    paragrafi_temizle(paragraph)

    satirlar = str(text or "").split("\n")
    if not satirlar:
        satirlar = [""]

    for i, satir in enumerate(satirlar):
        if i > 0:
            br = paragraph.add_run()
            format_uygula(br, fmt, sariyi_temizle=sariyi_temizle)
            br.add_break()
        run = paragraph.add_run(satir)
        format_uygula(run, fmt, sariyi_temizle=sariyi_temizle)


def vergi_dairesi_word_olustur(sablon_yolu: Path, cikti_yolu: Path, veri: Dict[str, str]) -> None:
    doc = Document(sablon_yolu)

    # Hitap alanı: genellikle ilk sarı paragraf bu alandır.
    sarili_paragraflar = sari_paragraflari_bul(doc)
    if not sarili_paragraflar:
        raise ValueError("Vergi Dairesi şablonunda sarı alan bulunamadı.")

    paragraf_yaz(sarili_paragraflar[0], vergi_dairesi_hitap_formatla(veri["vergi_dairesi_adi"]), fmt=paragraf_ilk_run_format(sarili_paragraflar[0]), sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
    kisi_ceza_ifadesini_guncelle(doc, veri)
    son_paragraf_ceza_ifadesini_guncelle(doc, veri)

    # Sayi / tarih alanları statik kalır.
    for p in belgedeki_paragraflar(doc):
        if "[SAYI]" in p.text or "[TARIH]" in p.text:
            yeni_metin = p.text
            paragraf_yaz(p, yeni_metin, fmt=paragraf_ilk_run_format(p), sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)

    # Konu alanı
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(text == "Konu" for text in row_text):
                hedef_cell = row.cells[2] if len(row.cells) > 2 else None
                if hedef_cell is None:
                    raise ValueError("Vergi Dairesi şablonunda konu hücresi bulunamadı.")
                hucre_yaz(hedef_cell, konu_metni_olustur(veri), font_size=Pt(12))
                break
        else:
            continue
        break

    # Dinamik tablo alanları
    target_table = None
    for table in doc.tables:
        if len(table.columns) >= 6 and len(table.rows) >= 2:
            header = [cell.text.strip() for cell in table.rows[0].cells]
            if "TCKN/VKN" in header and "SPK" in header[2]:
                target_table = table
                break

    if target_table is None:
        raise ValueError("Vergi Dairesi şablonunda dinamik kayıt tablosu bulunamadı.")

    tablo_kayitlari = belge_tablo_kayitlari(veri)
    tablo_satiri_yaz(target_table.rows[1], tablo_kayitlari[0])
    for tablo_veri in tablo_kayitlari[1:]:
        tablo_satiri_yaz(target_table.add_row(), tablo_veri)
    ayni_hucreleri_dikey_birlestir(target_table)
    tablo_tutarsizlik_yorumlarini_ekle(doc, target_table, tablo_kayitlari)

    if FINALDE_SARI_ALANLARI_TEMIZLE:
        tum_sari_vurgulari_temizle(doc)

    cikti_yolu.parent.mkdir(parents=True, exist_ok=True)
    doc.save(cikti_yolu)


def excelden_vergi_dairesi_kaydi_bul(excel_yolu: Path, kimlik_no: str, filtre: Optional[Dict[str, str]] = None) -> Dict[str, str]:
    wb = load_workbook(excel_yolu, data_only=True)
    if VERGI_DAIRESI_SHEET_ADI not in wb.sheetnames:
        raise ValueError(
            f"Excel dosyasında '{VERGI_DAIRESI_SHEET_ADI}' adlı sayfa bulunamadı. Mevcut sayfalar: {', '.join(wb.sheetnames)}"
        )

    ws = wb[VERGI_DAIRESI_SHEET_ADI]
    filtre = filtre or {}

    # İlk satır başlık, 2. satırdan başla
    for row in range(2, ws.max_row + 1):
        raw_kisi = hucre_metni(ws.cell(row=row, column=3))  # İlgili Kişi/Şirket column
        if not raw_kisi:
            continue
        if not satir_filtreye_uygun_mu(ws, row, filtre):
            continue

        kimlik_no_alan = vergi_dairesi_id_ayikla(raw_kisi)
        if not kimlik_no_alan:
            continue

        if kimlik_no_alan == kimlik_no:
            muhatap = vergi_dairesi_muhatap_ayikla(raw_kisi)
            vergi_dairesi_adi = hucre_metni(ws.cell(row=row, column=13))  # VD Adı column

            if not vergi_dairesi_adi:
                vergi_dairesi_adi = hucre_metni(ws.cell(row=row, column=2))
                if not vergi_dairesi_adi:
                    raise ValueError(f"Satır {row}: Vergi dairesi adı boş")

            detay_id = hucre_metni(ws.cell(row=row, column=10))
            tahsil_tarihi = hucre_metni(ws.cell(row=row, column=16), tarih_mi=True)
            ths_fis_no = hucre_metni(ws.cell(row=row, column=15))
            odeme_tutari = hucre_metni(ws.cell(row=row, column=17))

            if not vergi_dairesi_adi:
                vergi_dairesi_adi = hucre_metni(ws.cell(row=row, column=2))
            veri = spk_satir_verisi(ws, row, muhatap, kimlik_no, vergi_dairesi_adi)
            veri["parcali_tahsilatlar"] = spk_parcali_tahsilatlari_bul(ws, veri, row)
            return veri

    raise ValueError(f"'{kimlik_no}' TCKN/VKN ile kayıt bulunamadı")


def excelden_vergi_dairesi_kaydi_bul_girdi(excel_yolu: Path, girdi: str) -> Dict[str, str]:
    girdi_bilgisi = girdi_bilgisi_ayikla(girdi)
    temiz_girdi = girdi_bilgisi["arama"]
    rakamlar = re.sub(r"\D", "", temiz_girdi)
    if rakamlar and len(rakamlar) in (10, 11):
        return excelden_vergi_dairesi_kaydi_bul(excel_yolu, normalize_tckn(temiz_girdi), girdi_bilgisi)

    wb = load_workbook(excel_yolu, data_only=True)
    if VERGI_DAIRESI_SHEET_ADI not in wb.sheetnames:
        raise ValueError(
            f"Excel dosyasÄ±nda '{VERGI_DAIRESI_SHEET_ADI}' adlÄ± sayfa bulunamadÄ±. Mevcut sayfalar: {', '.join(wb.sheetnames)}"
        )

    ws = wb[VERGI_DAIRESI_SHEET_ADI]
    aranan = arama_metni_normalize(vergi_dairesi_muhatap_ayikla(temiz_girdi))

    for row in range(2, ws.max_row + 1):
        raw_kisi = hucre_metni(ws.cell(row=row, column=3))
        if not raw_kisi:
            continue
        if not satir_filtreye_uygun_mu(ws, row, girdi_bilgisi):
            continue

        muhatap = vergi_dairesi_muhatap_ayikla(raw_kisi)
        if arama_metni_normalize(muhatap) != aranan:
            continue

        vergi_dairesi_adi = hucre_metni(ws.cell(row=row, column=13))
        if not vergi_dairesi_adi:
            vergi_dairesi_adi = hucre_metni(ws.cell(row=row, column=2))
            if not vergi_dairesi_adi:
                raise ValueError(f"Satir {row}: Vergi dairesi adi bos")

        detay_id = hucre_metni(ws.cell(row=row, column=10))
        tahsil_tarihi = hucre_metni(ws.cell(row=row, column=16), tarih_mi=True)
        ths_fis_no = hucre_metni(ws.cell(row=row, column=15))
        odeme_tutari = hucre_metni(ws.cell(row=row, column=17))
        kimlik_no_alan = vergi_dairesi_id_ayikla(raw_kisi) or (rakamlar if len(rakamlar) in (10, 11) else "")

        veri = spk_satir_verisi(ws, row, muhatap, kimlik_no_alan, vergi_dairesi_adi)
        veri["parcali_tahsilatlar"] = spk_parcali_tahsilatlari_bul(ws, veri, row)
        return veri

    raise ValueError(f"'{temiz_girdi}' ile SPKYTMIPC kaydi bulunamadi")


def vergi_dairesi_turu_sec(root: tk.Tk) -> Optional[str]:
    sonuc = {"secim": None}

    pencere = tk.Toplevel(root)
    pencere.title("Vergi Dairesi Türü Seçimi")
    pencere.resizable(False, False)
    pencere.grab_set()

    etiket = tk.Label(
        pencere,
        text="Vergi Dairesi yazısı türünü seçiniz:",
        font=("Segoe UI", 11),
        padx=24,
        pady=18,
    )
    etiket.pack()

    buton_frame = tk.Frame(pencere, padx=18, pady=12)
    buton_frame.pack(fill="x")

    def sec(deger: str):
        sonuc["secim"] = deger
        pencere.destroy()

    btn_spk = tk.Button(
        buton_frame,
        text="SPK Bilgileri ile Yazı",
        width=24,
        height=2,
        command=lambda: sec("VERGI_DAIRESI_SPK"),
    )
    btn_spk.grid(row=0, column=0, padx=8, pady=6)

    btn_vd = tk.Button(
        buton_frame,
        text="Vergi Dairesi Bilgileri ile Yazı",
        width=24,
        height=2,
        command=lambda: sec("VERGI_DAIRESI_VD"),
    )
    btn_vd.grid(row=0, column=1, padx=8, pady=6)

    copyright_etiketi_ekle(pencere)

    def iptal():
        sonuc["secim"] = None
        pencere.destroy()

    pencere.protocol("WM_DELETE_WINDOW", iptal)

    pencere.update_idletasks()
    x = pencere.winfo_screenwidth() // 2 - pencere.winfo_width() // 2
    y = pencere.winfo_screenheight() // 2 - pencere.winfo_height() // 2
    pencere.geometry(f"+{x}+{y}")

    root.wait_window(pencere)
    return sonuc["secim"]


def vergi_dairesi_word_olustur(sablon_yolu: Path, cikti_yolu: Path, veri: Dict[str, str]) -> None:
    doc = Document(sablon_yolu)

    # Hitap alanı: genellikle ilk sarı paragraf bu alandır.
    sarili_paragraflar = sari_paragraflari_bul(doc)
    if not sarili_paragraflar:
        raise ValueError("Vergi Dairesi şablonunda sarı alan bulunamadı.")

    paragraf_yaz(sarili_paragraflar[0], vergi_dairesi_hitap_formatla(veri["vergi_dairesi_adi"]), fmt=paragraf_ilk_run_format(sarili_paragraflar[0]), sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
    kisi_ceza_ifadesini_guncelle(doc, veri)
    son_paragraf_ceza_ifadesini_guncelle(doc, veri)

    # Sayi / tarih alanları statik kalır.
    for p in belgedeki_paragraflar(doc):
        if "[SAYI]" in p.text or "[TARIH]" in p.text:
            yeni_metin = p.text
            paragraf_yaz(p, yeni_metin, fmt=paragraf_ilk_run_format(p), sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)

    # Konu alanı
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(text == "Konu" for text in row_text):
                hedef_cell = row.cells[2] if len(row.cells) > 2 else None
                if hedef_cell is None:
                    raise ValueError("Vergi Dairesi şablonunda konu hücresi bulunamadı.")
                hucre_yaz(hedef_cell, konu_metni_olustur(veri), font_size=Pt(12))
                break
        else:
            continue
        break

    # Dinamik tablo alanları
    target_table = None
    for table in doc.tables:
        if len(table.columns) >= 6 and len(table.rows) >= 2:
            header = [cell.text.strip() for cell in table.rows[0].cells]
            if "TCKN/VKN" in header and "SPK" in header[2]:
                target_table = table
                break

    if target_table is None:
        raise ValueError("Vergi Dairesi şablonunda dinamik kayıt tablosu bulunamadı.")

    tablo_kayitlari = belge_tablo_kayitlari(veri)
    tablo_satiri_yaz(target_table.rows[1], tablo_kayitlari[0])
    for tablo_veri in tablo_kayitlari[1:]:
        tablo_satiri_yaz(target_table.add_row(), tablo_veri)
    ayni_hucreleri_dikey_birlestir(target_table)
    tablo_tutarsizlik_yorumlarini_ekle(doc, target_table, tablo_kayitlari)

    if FINALDE_SARI_ALANLARI_TEMIZLE:
        tum_sari_vurgulari_temizle(doc)

    cikti_yolu.parent.mkdir(parents=True, exist_ok=True)
    doc.save(cikti_yolu)


# =========================
# YARDIMCI FONKSİYONLAR
# =========================

def uygulama_klasoru() -> Path:
    """Script hangi klasördeyse onu döndürür. PyInstaller ile paketlenirse exe klasörünü kullanır."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def kaynak_klasoru() -> Path:
    """PyInstaller tek dosya paketinde gomulu kaynaklarin acildigi klasoru dondurur."""
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS)
    return Path(__file__).resolve().parent


def temiz_dosya_adi(metin: str) -> str:
    metin = str(metin or "").strip()
    metin = re.sub(r"[\\/:*?\"<>|]+", "_", metin)
    metin = re.sub(r"\s+", "_", metin)
    return metin.strip("_") or "belge"


def normalize_tckn(value) -> str:
    """Excel/Tkinter girdisinden 11 haneli TCKN metni üretir."""
    if value is None:
        return ""

    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return ""

    text = str(value).strip()
    digits = re.sub(r"\D", "", text)

    if len(digits) == 10:
        return digits.zfill(11)
    if len(digits) == 11:
        return digits
    if len(digits) < 10:
        return digits.zfill(10)
    return digits


def kullanici_kimlik_no_denetle(value: str) -> Tuple[Optional[str], Optional[str]]:
    text = str(value or "").strip()
    digits = re.sub(r"\D", "", text)

    if not digits:
        return None, "rakam içermiyor"
    if len(digits) == 10:
        return digits, None
    if len(digits) == 11:
        if int(digits[-1]) % 2 != 0:
            return None, "TCKN son hanesi çift sayı olmalı"
        return digits, None
    return None, "VKN 10 hane, TCKN 11 hane olmalı"


def hucre_metni(cell, tarih_mi: bool = False) -> str:
    """Openpyxl hücresini temiz metne çevirir."""
    return excel_degeri_metni(cell.value, tarih_mi=tarih_mi)


def excel_degeri_metni(value, tarih_mi: bool = False) -> str:
    """Openpyxl hucre degerini temiz metne cevirir."""
    if value is None:
        return ""

    if tarih_mi:
        if isinstance(value, (datetime, date)):
            return value.strftime("%d.%m.%Y")
        # Hücre tarih formatlı ise openpyxl genelde datetime döndürür; metinse olduğu gibi korunur.
        return str(value).strip()

    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value).strip()

    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")

    return str(value).strip()


def tahsilat_tarihi_fis_no_metni(tahsil_tarihi: str, ths_fis_no: str) -> str:
    parcalar = [str(parca or "").strip() for parca in (tahsil_tarihi, ths_fis_no)]
    return " - ".join(parca for parca in parcalar if parca)


def spk_satir_verisi(ws, row: int, muhatap: str, kimlik_no: str, vergi_dairesi_adi: str) -> Dict[str, str]:
    ipc_dosya_no = hucre_kodu_metni(ws.cell(row=row, column=1).value)
    daire = hucre_metni(ws.cell(row=row, column=2))
    detay_id = hucre_metni(ws.cell(row=row, column=10))
    vd_kodu = hucre_kodu_metni(ws.cell(row=row, column=12).value)
    tahsil_tarihi = hucre_metni(ws.cell(row=row, column=16), tarih_mi=True)
    ths_fis_no = hucre_metni(ws.cell(row=row, column=15))
    odeme_tutari = hucre_metni(ws.cell(row=row, column=17))
    return {
        "muhatap": muhatap,
        "kimlik_no": kimlik_no,
        "vergi_dairesi_adi": vergi_dairesi_adi,
        "karar_tarihi": hucre_metni(ws.cell(row=row, column=4), tarih_mi=True),
        "karar_no": hucre_metni(ws.cell(row=row, column=5)),
        "pc_tutari": odeme_tutari,
        "odenen_tutar": odeme_tutari,
        "tahsilat_tarihi_fis_no": tahsilat_tarihi_fis_no_metni(tahsil_tarihi, ths_fis_no),
        "ipc_dosya_no": ipc_dosya_no,
        "daire": daire,
        "vd_kodu": vd_kodu,
        "sayi": detay_id,
        "tarih": tahsil_tarihi,
        "konu": konu_metni_olustur({"muhatap": muhatap, "kimlik_no": kimlik_no}),
        "excel_sayfa": VERGI_DAIRESI_SHEET_ADI,
        "excel_satir": str(row),
        "onceki_excel_satirlari": "",
    }


def spk_parcali_tahsilatlari_bul(ws, ana_veri: Dict[str, str], ana_satir: int) -> List[Dict[str, str]]:
    ana_ipc = ipc_tarih_sayi_anahtari(ana_veri.get("karar_tarihi", ""), ana_veri.get("karar_no", ""))
    ana_ipc_dosya_no = str(ana_veri.get("ipc_dosya_no", "")).strip()
    ana_daire = arama_metni_normalize(ana_veri.get("daire", ""))
    sonuc = []
    onceki_satirlar = []

    for row in range(2, ws.max_row + 1):
        if row == ana_satir:
            continue

        raw_kisi = hucre_metni(ws.cell(row=row, column=3))
        if not raw_kisi:
            continue

        kimlik_no_alan = vergi_dairesi_id_ayikla(raw_kisi)
        muhatap = vergi_dairesi_muhatap_ayikla(raw_kisi)
        ayni_kisi = kimlik_no_alan == ana_veri.get("kimlik_no") if kimlik_no_alan else arama_metni_normalize(muhatap) == arama_metni_normalize(ana_veri.get("muhatap", ""))
        if not ayni_kisi:
            continue

        if ana_ipc_dosya_no:
            satir_ipc_dosya_no = hucre_kodu_metni(ws.cell(row=row, column=1).value)
            if satir_ipc_dosya_no != ana_ipc_dosya_no:
                continue
            if ana_daire and arama_metni_normalize(hucre_metni(ws.cell(row=row, column=2))) != ana_daire:
                continue
        else:
            satir_ipc = ipc_tarih_sayi_anahtari(hucre_metni(ws.cell(row=row, column=4), tarih_mi=True), hucre_metni(ws.cell(row=row, column=5)))
            if satir_ipc != ana_ipc:
                continue

        parcali = spk_satir_verisi(
            ws,
            row,
            ana_veri["muhatap"],
            ana_veri["kimlik_no"],
            ana_veri["vergi_dairesi_adi"],
        )
        sonuc.append(parcali)
        onceki_satirlar.append(str(row))

    ana_veri["onceki_excel_satirlari"] = ", ".join(onceki_satirlar)
    return sonuc


def odenen_tutari_word_icin_formatla(tutar: str) -> str:
    """Word icin 1.234,50 biciminde para metni uretir."""
    sayi = para_decimal(tutar)
    if sayi is not None:
        return decimal_para_metni(sayi)

    metin = str(tutar or "").strip().replace("₺", "").replace("TL", "").strip()
    if not metin:
        return ""

    temiz = re.sub(r"[^0-9,.\-]", "", metin)
    if not temiz or temiz in {"-", ".", ","}:
        return metin

    son_nokta = temiz.rfind(".")
    son_virgul = temiz.rfind(",")
    decimal_ayiraci = ""

    if son_nokta != -1 and son_virgul != -1:
        decimal_ayiraci = "." if son_nokta > son_virgul else ","
    elif son_nokta != -1:
        decimal_hane = len(temiz) - son_nokta - 1
        if decimal_hane in (1, 2):
            decimal_ayiraci = "."
    elif son_virgul != -1:
        decimal_hane = len(temiz) - son_virgul - 1
        if decimal_hane in (1, 2):
            decimal_ayiraci = ","

    if decimal_ayiraci:
        tam, ondalik = temiz.rsplit(decimal_ayiraci, 1)
        tam = re.sub(r"[,.]", "", tam)
        sayi_metni = f"{tam}.{ondalik}"
    else:
        sayi_metni = re.sub(r"[,.]", "", temiz)

    try:
        sayi = Decimal(sayi_metni).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    except InvalidOperation:
        return metin

    return decimal_para_metni(sayi)


def para_decimal(tutar: str) -> Optional[Decimal]:
    metin = str(tutar or "").strip().replace("₺", "").replace("TL", "").strip()
    if not metin:
        return None

    temiz = re.sub(r"[^0-9,.\-]", "", metin)
    if not temiz or temiz in {"-", ".", ","}:
        return None

    son_nokta = temiz.rfind(".")
    son_virgul = temiz.rfind(",")
    decimal_ayiraci = ""

    if son_nokta != -1 and son_virgul != -1:
        decimal_ayiraci = "." if son_nokta > son_virgul else ","
    elif son_nokta != -1 and len(temiz) - son_nokta - 1 in (1, 2):
        decimal_ayiraci = "."
    elif son_virgul != -1 and len(temiz) - son_virgul - 1 in (1, 2):
        decimal_ayiraci = ","

    if decimal_ayiraci:
        tam, ondalik = temiz.rsplit(decimal_ayiraci, 1)
        sayi_metni = f"{re.sub(r'[,.]', '', tam)}.{ondalik}"
    else:
        sayi_metni = re.sub(r"[,.]", "", temiz)

    try:
        return Decimal(sayi_metni).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    except InvalidOperation:
        return None


def decimal_para_metni(sayi: Decimal) -> str:
    sayi = sayi.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"{sayi:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def toplam_odenen_tutar_decimal(veri: Dict[str, Any]) -> Optional[Decimal]:
    toplam = para_decimal(veri.get("odenen_tutar", ""))
    if toplam is None:
        return None
    for parcali in veri.get("parcali_tahsilatlar", []):
        tutar = para_decimal(parcali.get("odenen_tutar", ""))
        if tutar is not None:
            toplam += tutar
    return toplam.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def para_metni_ikiyle_carp(tutar: str) -> str:
    sayi = para_decimal(tutar)
    if sayi is None:
        return tutar
    return decimal_para_metni(sayi * Decimal("2"))


def arama_metni_normalize(metin: str) -> str:
    metin = turkce_buyuk_harf(str(metin or "").strip())
    metin = re.sub(r"[^0-9A-ZÇĞİÖŞÜ]+", " ", metin)
    return re.sub(r"\s+", " ", metin).strip()


def borclu_bilgisi_unvan_ayikla(metin: str) -> str:
    metin = str(metin or "").strip()
    metin = re.sub(r"\s+[İI]PC\s+Tahakkuk\s*\([^)]*\)\s*$", "", metin, flags=re.IGNORECASE).strip()
    return re.sub(r"\s*\(\d+\)\s*$", "", metin).strip()


def ipc_tarih_sayi_anahtari(tarih: str, sayi: str) -> str:
    metin = f"{tarih}-{sayi}"
    metin = re.sub(r"\s+", "", str(metin or ""))
    metin = metin.replace("\\", "/")
    return metin.upper()


def ipc_tarih_sayi_metni_normalize(metin: str) -> str:
    metin = str(metin or "").strip()
    eslesme = re.search(r"\((\d{1,2}\.\d{1,2}\.\d{4}\s*-\s*[^)]+)\)", metin)
    if eslesme:
        metin = eslesme.group(1)
    metin = re.sub(r"\s+", "", metin)
    metin = metin.replace("\\", "/")
    return metin.upper()


def tarih_metni_parse(metin: str) -> Optional[date]:
    metin = str(metin or "").strip()
    for fmt in ("%d.%m.%Y", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(metin, fmt).date()
        except ValueError:
            continue
    return None


def ipc_anahtarindan_tarih(ipc: str) -> Optional[date]:
    eslesme = re.match(r"(\d{1,2}\.\d{1,2}\.\d{4})-", str(ipc or ""))
    return tarih_metni_parse(eslesme.group(1)) if eslesme else None


def para_metni_word_icin_formatla(tutar: str) -> str:
    sayi = para_decimal(tutar)
    if sayi is not None:
        return decimal_para_metni(sayi)
    return str(tutar or "").strip().replace("₺", "").strip()


def kalan_alacak_kaydi_indekse_ekle(indeks: Dict[str, Any], tur: str, kisi_anahtari: str, kayit: Dict[str, Any]) -> None:
    ipc = kayit["ipc"]
    indeks[f"{tur}_kisi"].add(kisi_anahtari)
    anahtar = f"{kisi_anahtari}|{ipc}"
    mevcut = indeks[tur].get(anahtar)

    if mevcut:
        mevcut_tutar = mevcut.get("pc_tutari_decimal") or Decimal("0.00")
        yeni_tutar = kayit.get("pc_tutari_decimal") or Decimal("0.00")
        toplam = (mevcut_tutar + yeni_tutar).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        mevcut["pc_tutari_decimal"] = toplam
        mevcut["pc_tutari"] = decimal_para_metni(toplam)
        mevcut.setdefault("rows", [mevcut.get("row")]).append(kayit.get("row"))
        return

    kayit = dict(kayit)
    kayit["rows"] = [kayit.get("row")]
    indeks[tur][anahtar] = kayit
    indeks[f"{tur}_kayitlari"].setdefault(kisi_anahtari, []).append(kayit)


def kalan_alacak_indeksi_olustur(kalan_alacak_yolu: Path) -> Dict[str, Any]:
    wb = load_workbook(kalan_alacak_yolu, data_only=True, read_only=True)
    ws = wb[wb.sheetnames[0]]
    indeks: Dict[str, Any] = {
        "tckn": {},
        "unvan": {},
        "tckn_kayitlari": {},
        "unvan_kayitlari": {},
        "tckn_kisi": set(),
        "unvan_kisi": set(),
    }

    for row_no, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        ipc = ipc_tarih_sayi_metni_normalize(excel_degeri_metni(row[4] if len(row) > 4 else ""))
        borclu_bilgisi = excel_degeri_metni(row[16] if len(row) > 16 else "")
        kaynak_tutar = para_metni_word_icin_formatla(excel_degeri_metni(row[19] if len(row) > 19 else ""))
        tutar = para_metni_ikiyle_carp(kaynak_tutar)
        if not borclu_bilgisi or not ipc:
            continue
        kayit = {
            "row": row_no,
            "ipc": ipc,
            "karar_tarihi": ipc.split("-", 1)[0] if "-" in ipc else "",
            "karar_no": ipc.split("-", 1)[1] if "-" in ipc else "",
            "pc_tutari": tutar,
            "pc_tutari_decimal": para_decimal(tutar),
            "kaynak_pc_tutari": kaynak_tutar,
            "borclu_bilgisi": borclu_bilgisi,
        }

        for kimlik in re.findall(r"\((\d{10,11})\)", borclu_bilgisi):
            kalan_alacak_kaydi_indekse_ekle(indeks, "tckn", kimlik, kayit)

        borclu_unvan = arama_metni_normalize(borclu_bilgisi_unvan_ayikla(borclu_bilgisi))
        if borclu_unvan:
            kalan_alacak_kaydi_indekse_ekle(indeks, "unvan", borclu_unvan, kayit)

    return indeks


def kalan_alacak_unvan_arama_kelimeleri(unvan: str) -> List[str]:
    haric = {
        "A", "S", "AS", "AŞ", "ANONIM", "ANONİM", "SIRKET", "ŞIRKET", "ŞİRKET",
        "SIRKETI", "ŞIRKETI", "ŞİRKETİ", "LIMITED", "LTD", "STI", "ŞTI", "ŞTİ",
        "TICARET", "TİCARET",
    }
    kelimeler = arama_metni_normalize(unvan).split()
    anlamli = [kelime for kelime in kelimeler if len(kelime) > 1 and not kelime.isdigit() and kelime not in haric]
    return anlamli or [kelime for kelime in kelimeler if kelime]


def kalan_alacak_unvan_anahtari(veri: Dict[str, str], girdi: str) -> str:
    return str(veri.get("kalan_alacak_unvan_anahtari") or arama_metni_normalize(girdi or veri.get("muhatap", "")))


def kalan_alacak_muhtemel_unvan_eslesmeleri(kalan_alacak_indeksi: Dict[str, Any], unvan: str) -> List[Dict[str, Any]]:
    kelimeler = kalan_alacak_unvan_arama_kelimeleri(unvan)
    if not kelimeler:
        return []

    adaylar = []
    for unvan_anahtari, kayitlar in kalan_alacak_indeksi["unvan_kayitlari"].items():
        unvan_kelimeleri = set(str(unvan_anahtari).split())
        if not all(kelime in unvan_kelimeleri for kelime in kelimeler):
            continue

        satirlar = sorted({k.get("row") for k in kayitlar if k.get("row")})
        adaylar.append(
            {
                "unvan_anahtari": unvan_anahtari,
                "kayit_sayisi": len(kayitlar),
                "satirlar": satirlar,
                "ilk_kayit": kayitlar[0] if kayitlar else {},
                "vkn": next(
                    (
                        kimlik
                        for kayit in kayitlar
                        for kimlik in re.findall(r"\((\d{10})\)", str(kayit.get("borclu_bilgisi", "")))
                    ),
                    "",
                ),
                "tam_eslesme": unvan_anahtari == arama_metni_normalize(unvan),
                "kelime_sayisi": len(unvan_kelimeleri),
            }
        )

    return sorted(adaylar, key=lambda a: (not a["tam_eslesme"], a["kelime_sayisi"], a["unvan_anahtari"]))[:25]


def kalan_alacak_unvan_vkn_bul(kalan_alacak_indeksi: Dict[str, Any], unvan_anahtari: str) -> str:
    for kayit in kalan_alacak_indeksi["unvan_kayitlari"].get(unvan_anahtari, []):
        for kimlik in re.findall(r"\((\d{10,11})\)", str(kayit.get("borclu_bilgisi", ""))):
            if len(kimlik) == 10:
                return kimlik
    return ""


def kalan_alacak_unvan_eslesmesini_sec(root, girdi: str, adaylar: List[Dict[str, Any]]) -> Optional[str]:
    if not adaylar:
        return None

    sonuc = {"secim": None}
    pencere = tk.Toplevel(root)
    pencere.title("Kalan Alacak Ünvan Eşleşmesi")
    pencere.resizable(True, True)
    pencere.grab_set()

    aciklama = tk.Label(
        pencere,
        text=f"'{girdi}' için Kalan Alacak dosyasında bulunan muhtemel ünvan eşleşmesini seçiniz.",
        font=("Segoe UI", 10),
        justify="left",
        padx=16,
        pady=12,
    )
    aciklama.pack(fill="x")

    liste = tk.Listbox(pencere, selectmode=tk.SINGLE, width=120, height=min(12, max(4, len(adaylar))), font=("Consolas", 9))
    liste.pack(fill="both", expand=True, padx=16, pady=(0, 10))

    for aday in adaylar:
        satirlar = ", ".join(str(satir) for satir in aday["satirlar"][:5])
        if len(aday["satirlar"]) > 5:
            satirlar += ", ..."
        vkn = f" | VKN: {aday['vkn']}" if aday.get("vkn") else ""
        liste.insert(
            tk.END,
            f"{aday['unvan_anahtari']}{vkn} | Kayıt: {aday['kayit_sayisi']} | Satır: {satirlar}",
        )
    liste.selection_set(0)

    buton_frame = tk.Frame(pencere, padx=16, pady=12)
    buton_frame.pack(fill="x")

    def sec():
        secimler = list(liste.curselection())
        sonuc["secim"] = adaylar[secimler[0]]["unvan_anahtari"] if secimler else None
        pencere.destroy()

    def atla():
        sonuc["secim"] = None
        pencere.destroy()

    tk.Button(buton_frame, text="Seç", width=12, command=sec).pack(side="right", padx=(6, 0))
    tk.Button(buton_frame, text="Atla", width=12, command=atla).pack(side="right")

    pencere.protocol("WM_DELETE_WINDOW", atla)
    pencere.update_idletasks()
    x = pencere.winfo_screenwidth() // 2 - pencere.winfo_width() // 2
    y = pencere.winfo_screenheight() // 2 - pencere.winfo_height() // 2
    pencere.geometry(f"+{x}+{y}")
    root.wait_window(pencere)
    return sonuc["secim"]


def kalan_alacak_unvan_eslesmesini_veriye_uygula(root, kalan_alacak_indeksi: Dict[str, Any], girdi: str, veri: Dict[str, str]) -> Optional[str]:
    if kimlik_turu_belirle(veri.get("kimlik_no", "")) == "TCKN":
        return None

    girdi_bilgisi = girdi_bilgisi_ayikla(girdi)
    rakamlar = re.sub(r"\D", "", girdi_bilgisi["arama"])
    arama_unvani = veri.get("muhatap", "") if rakamlar == girdi_bilgisi["arama"].strip() else vergi_dairesi_muhatap_ayikla(girdi_bilgisi["arama"])
    adaylar = kalan_alacak_muhtemel_unvan_eslesmeleri(kalan_alacak_indeksi, arama_unvani)
    secilen = kalan_alacak_unvan_eslesmesini_sec(root, arama_unvani, adaylar)
    if secilen:
        veri["kalan_alacak_unvan_anahtari"] = secilen
        vkn = kalan_alacak_unvan_vkn_bul(kalan_alacak_indeksi, secilen)
        if vkn:
            veri["kimlik_no"] = vkn
    return secilen


def kalan_alacak_ipc_tutari_bul(kalan_alacak_indeksi: Dict[str, Any], girdi: str, veri: Dict[str, str]) -> Tuple[Optional[str], str]:
    arama_girdisi = girdi_bilgisi_ayikla(girdi)["arama"]
    rakamlar = re.sub(r"\D", "", arama_girdisi)
    tckn = normalize_tckn(arama_girdisi) if rakamlar and len(rakamlar) == 11 else ""
    aranan_unvan = kalan_alacak_unvan_anahtari(veri, arama_girdisi if not tckn else veri.get("muhatap", ""))
    aranan_ipc = ipc_tarih_sayi_anahtari(veri.get("karar_tarihi", ""), veri.get("karar_no", ""))

    if tckn:
        sonuc = kalan_alacak_indeksi["tckn"].get(f"{tckn}|{aranan_ipc}")
        kisi_bulundu = tckn in kalan_alacak_indeksi["tckn_kisi"]
    else:
        sonuc = kalan_alacak_indeksi["unvan"].get(f"{aranan_unvan}|{aranan_ipc}")
        kisi_bulundu = aranan_unvan in kalan_alacak_indeksi["unvan_kisi"]

    if sonuc:
        tutar, row = sonuc["pc_tutari"], sonuc["row"]
        return tutar, f"{arama_girdisi}: Kalan Alacak satir {row} bulundu, IPC {aranan_ipc} eslesti, IPC Tutari {tutar} yazildi."

    if kisi_bulundu:
        return None, f"{arama_girdisi}: Kalan Alacak dosyasinda kisi bulundu ancak IPC tarih/sayi ({aranan_ipc}) eslesmedi; IPC Tutari mevcut Excel tutariyla kaldi."
    return None, f"{arama_girdisi}: Kalan Alacak dosyasinda eslesme bulunamadi; IPC Tutari mevcut Excel tutariyla kaldi."


def ipc_beklenen_odeme_tutari(kayit: Dict[str, Any], tahsil_tarihi: Optional[date]) -> Optional[Decimal]:
    tutar = kayit.get("pc_tutari_decimal")
    if tutar is None:
        return None

    ipc_tarihi = ipc_anahtarindan_tarih(kayit.get("ipc", ""))
    if tahsil_tarihi and ipc_tarihi:
        gun_farki = (tahsil_tarihi - ipc_tarihi).days
        if 0 <= gun_farki < 60:
            return (tutar * Decimal("0.75")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return tutar


def kisi_ipc_kayitlari(kalan_alacak_indeksi: Dict[str, Any], girdi: str, veri: Dict[str, str]) -> List[Dict[str, Any]]:
    arama_girdisi = girdi_bilgisi_ayikla(girdi)["arama"]
    rakamlar = re.sub(r"\D", "", arama_girdisi)
    tckn = normalize_tckn(arama_girdisi) if rakamlar and len(rakamlar) == 11 else ""
    aranan_unvan = kalan_alacak_unvan_anahtari(veri, arama_girdisi if not tckn else veri.get("muhatap", ""))
    return kalan_alacak_indeksi["tckn_kayitlari"].get(tckn, []) if tckn else kalan_alacak_indeksi["unvan_kayitlari"].get(aranan_unvan, [])


def ipc_adaylarini_hazirla(kalan_alacak_indeksi: Dict[str, Any], girdi: str, veri: Dict[str, str]) -> Tuple[Optional[Dict[str, Any]], List[Dict[str, Any]]]:
    ana_ipc = ipc_tarih_sayi_anahtari(veri.get("karar_tarihi", ""), veri.get("karar_no", ""))
    tahsil_tarihi = tarih_metni_parse(veri.get("tarih", ""))
    ana_kayit = None
    adaylar = []
    gorulen = set()

    for kayit in kisi_ipc_kayitlari(kalan_alacak_indeksi, girdi, veri):
        satir = kayit.get("row")
        if satir in gorulen:
            continue
        gorulen.add(satir)

        beklenen = ipc_beklenen_odeme_tutari(kayit, tahsil_tarihi)
        if beklenen is None:
            continue
        kayit = dict(kayit)
        kayit["beklenen_odeme"] = beklenen
        if kayit["ipc"] == ana_ipc:
            if ana_kayit is None:
                ana_kayit = kayit
            continue
        else:
            adaylar.append(kayit)

    return ana_kayit, adaylar


def odemeyi_aciklayan_ipcleri_bul(kalan_alacak_indeksi: Dict[str, Any], girdi: str, veri: Dict[str, str]) -> List[Dict[str, Any]]:
    odenen = toplam_odenen_tutar_decimal(veri)
    ana_tutar = para_decimal(veri.get("pc_tutari", ""))
    if odenen is None or ana_tutar is None or ana_tutar >= odenen:
        return []

    ana_ipc = ipc_tarih_sayi_anahtari(veri.get("karar_tarihi", ""), veri.get("karar_no", ""))
    ana_kayit, adaylar = ipc_adaylarini_hazirla(kalan_alacak_indeksi, girdi, veri)

    if ana_kayit is None:
        ana_kayit = {
            "ipc": ana_ipc,
            "pc_tutari": veri.get("pc_tutari", ""),
            "pc_tutari_decimal": ana_tutar,
            "beklenen_odeme": ana_tutar,
            "row": "",
        }

    hedef = odenen
    baslangic = ana_kayit["beklenen_odeme"]
    kalan = (hedef - baslangic).copy_abs()
    if kalan <= Decimal("0.01"):
        return []

    adaylar = [a for a in adaylar if a["beklenen_odeme"] <= kalan + Decimal("0.01")]
    for adet in range(1, min(len(adaylar), 8) + 1):
        from itertools import combinations
        for kombinasyon in combinations(adaylar, adet):
            toplam = baslangic + sum((a["beklenen_odeme"] for a in kombinasyon), Decimal("0.00"))
            if abs(toplam - hedef) <= Decimal("0.01"):
                return list(kombinasyon)
    return []


def muhtemel_ipcleri_bul(kalan_alacak_indeksi: Dict[str, Any], girdi: str, veri: Dict[str, str]) -> List[Dict[str, Any]]:
    odenen = toplam_odenen_tutar_decimal(veri)
    ana_tutar = para_decimal(veri.get("pc_tutari", ""))
    if odenen is None or ana_tutar is None or ana_tutar >= odenen:
        return []

    ana_kayit, adaylar = ipc_adaylarini_hazirla(kalan_alacak_indeksi, girdi, veri)
    ana_beklenen = ana_kayit["beklenen_odeme"] if ana_kayit else ana_tutar
    kalan = odenen - ana_beklenen
    if kalan <= Decimal("0.01"):
        return []

    for kayit in adaylar:
        kayit["fark"] = abs(kalan - kayit["beklenen_odeme"])
        kayit["gun_farki"] = None
        ipc_tarihi = ipc_anahtarindan_tarih(kayit.get("ipc", ""))
        tahsil_tarihi = tarih_metni_parse(veri.get("tarih", ""))
        if ipc_tarihi and tahsil_tarihi:
            kayit["gun_farki"] = (tahsil_tarihi - ipc_tarihi).days

    return sorted(adaylar, key=lambda k: (k["fark"], k.get("row") or 0))[:10]


def ek_ipcleri_kullaniciya_sor(root, veri: Dict[str, str], ek_ipcler: List[Dict[str, Any]]) -> bool:
    if not ek_ipcler:
        return False
    toplam_odenen = toplam_odenen_tutar_decimal(veri)
    satirlar = [
        f"Ödenen tutar toplamı ({decimal_para_metni(toplam_odenen) if toplam_odenen is not None else odenen_tutari_word_icin_formatla(veri.get('odenen_tutar', ''))}) tek IPC tutarından büyüktür.",
        "Aşağıdaki aynı kişiye ait IPC kayıtları ödemeyi açıklıyor. Tabloya eklensin mi?",
        "",
    ]
    for kayit in ek_ipcler:
        satirlar.append(
            f"- {kayit['ipc']} | IPC Tutarı: {kayit['pc_tutari']} | Hesaba katılan: {decimal_para_metni(kayit['beklenen_odeme'])}"
        )
    return messagebox.askyesno("Ek IPC kayıtları", "\n".join(satirlar), parent=root)


def muhtemel_ipcleri_kullaniciya_sec(root, veri: Dict[str, str], adaylar: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not adaylar:
        return []

    sonuc = {"secim": []}
    toplam_odenen = toplam_odenen_tutar_decimal(veri)
    pencere = tk.Toplevel(root)
    pencere.title("Muhtemel IPC Seçimi")
    pencere.resizable(True, True)
    pencere.grab_set()

    aciklama = tk.Label(
        pencere,
        text=(
            f"Ödenen tutar toplamı ({decimal_para_metni(toplam_odenen) if toplam_odenen is not None else odenen_tutari_word_icin_formatla(veri.get('odenen_tutar', ''))}) tek IPC tutarından büyüktür.\n"
            "Muhtemel ek IPC kayıtlarını seçiniz."
        ),
        font=("Segoe UI", 10),
        justify="left",
        padx=16,
        pady=12,
    )
    aciklama.pack(fill="x")

    liste = tk.Listbox(pencere, selectmode=tk.MULTIPLE, width=110, height=min(10, max(4, len(adaylar))), font=("Consolas", 9))
    liste.pack(fill="both", expand=True, padx=16, pady=(0, 10))

    for kayit in adaylar:
        gun = "" if kayit.get("gun_farki") is None else f" | Gün: {kayit['gun_farki']}"
        liste.insert(
            tk.END,
            f"{kayit['ipc']} | IPC: {kayit['pc_tutari']} | Hesaba katılan: {decimal_para_metni(kayit['beklenen_odeme'])} | Fark: {decimal_para_metni(kayit['fark'])}{gun}",
        )

    if adaylar:
        liste.selection_set(0)

    buton_frame = tk.Frame(pencere, padx=16, pady=12)
    buton_frame.pack(fill="x")

    def ekle():
        secimler = list(liste.curselection())
        sonuc["secim"] = [adaylar[i] for i in secimler]
        pencere.destroy()

    def atla():
        sonuc["secim"] = []
        pencere.destroy()

    tk.Button(buton_frame, text="Seçilenleri Ekle", width=18, command=ekle).pack(side="right", padx=(6, 0))
    tk.Button(buton_frame, text="Ekleme", width=12, command=atla).pack(side="right")

    pencere.protocol("WM_DELETE_WINDOW", atla)
    pencere.update_idletasks()
    x = pencere.winfo_screenwidth() // 2 - pencere.winfo_width() // 2
    y = pencere.winfo_screenheight() // 2 - pencere.winfo_height() // 2
    pencere.geometry(f"+{x}+{y}")
    root.wait_window(pencere)
    return sonuc["secim"]


def turkce_buyuk_harf(metin: str) -> str:
    """Türkçe i/ı dönüşümlerini koruyarak büyük harfe çevirir."""
    return str(metin or "").translate(str.maketrans({"i": "İ", "ı": "I"})).upper()


def turkce_kucuk_harf(metin: str) -> str:
    """Türkçe I/İ dönüşümlerini koruyarak küçük harfe çevirir."""
    return str(metin or "").translate(str.maketrans({"I": "ı", "İ": "i"})).lower()


def turkce_bas_harf_buyut(metin: str) -> str:
    metin = turkce_kucuk_harf(metin)
    if not metin:
        return ""
    return turkce_buyuk_harf(metin[0]) + metin[1:]


def ad_soyad_formatla(metin: str) -> str:
    parcalar = re.split(r"\s+", str(metin or "").strip())
    if not parcalar:
        return ""
    if len(parcalar) == 1:
        return turkce_bas_harf_buyut(parcalar[0])

    adlar = [turkce_bas_harf_buyut(parca) for parca in parcalar[:-1]]
    soyad = turkce_buyuk_harf(parcalar[-1])
    return " ".join(adlar + [soyad])


def copyright_etiketi_ekle(pencere) -> None:
    """Pencere altına copyright etiketi ekler."""
    etiket = tk.Label(
        pencere,
        text=COPYRIGHT,
        font=("Segoe UI", 8),
        fg="gray",
    )
    etiket.pack(side="bottom", pady=(10, 0))


def dosya_yoksa_sec(root, varsayilan_yol: Path, baslik: str, dosya_turleri) -> Optional[Path]:
    """Dosya varsa varsayılan yolu kullan, yoksa seçtir."""
    if varsayilan_yol.exists():
        return varsayilan_yol
    return dosya_sec(root, baslik, dosya_turleri)


def dosya_sec(root, baslik: str, dosya_turleri) -> Optional[Path]:
    """Dosya seçme dialogu."""
    dosya_yolu = filedialog.askopenfilename(
        parent=root,
        title=baslik,
        filetypes=dosya_turleri,
    )
    return Path(dosya_yolu) if dosya_yolu else None


def tcknleri_al(root) -> Optional[List[str]]:
    """Kullanıcıdan TCKN/VKN listesi alır."""
    sonuc = {"tcknler": None}

    pencere = tk.Toplevel(root)
    pencere.title("TCKN/VKN Girişi")
    pencere.resizable(False, False)
    pencere.grab_set()

    etiket = tk.Label(
        pencere,
        text="Gerçek kişi için TCKN, tüzel kişi için ünvan giriniz. İsterseniz IPC id ve daire kodunu da ekleyiniz (her satıra bir tane):",
        font=("Segoe UI", 11),
        padx=24,
        pady=18,
    )
    etiket.pack()

    metin_alani = tk.Text(pencere, width=96, height=10, font=("Consolas", 10))
    metin_alani.pack(padx=24, pady=(0, 12))

    buton_frame = tk.Frame(pencere, padx=18, pady=12)
    buton_frame.pack(fill="x")

    def tamamla():
        metin = metin_alani.get("1.0", tk.END).strip()
        if not metin:
            messagebox.showerror("Hata", "En az bir TCKN veya ünvan giriniz.", parent=pencere)
            return

        satirlar = [s.strip() for s in metin.split("\n") if s.strip()]
        tcknler = []
        hatalar = []

        for satir in satirlar:
            rakamlar = re.sub(r"\D", "", satir)
            if rakamlar and rakamlar == satir.strip():
                temizlenmis, hata = kullanici_kimlik_no_denetle(satir)
                if hata:
                    hatalar.append(f"'{satir}': {hata}")
                elif temizlenmis:
                    tcknler.append(temizlenmis)
            else:
                tcknler.append(satir.strip())

        if hatalar:
            messagebox.showerror(
                "Hata",
                "Geçersiz TCKN/VKN:\n" + "\n".join(hatalar),
                parent=pencere,
            )
            return

        if not tcknler:
            messagebox.showerror("Hata", "Geçerli TCKN/VKN bulunamadı.", parent=pencere)
            return

        sonuc["tcknler"] = tcknler
        pencere.destroy()

    btn_tamam = tk.Button(
        buton_frame,
        text="Tamam",
        width=12,
        height=2,
        command=tamamla,
    )
    btn_tamam.pack(side="right", padx=(6, 0))

    btn_iptal = tk.Button(
        buton_frame,
        text="İptal",
        width=12,
        height=2,
        command=pencere.destroy,
    )
    btn_iptal.pack(side="right")

    copyright_etiketi_ekle(pencere)

    def iptal():
        sonuc["tcknler"] = None
        pencere.destroy()

    pencere.protocol("WM_DELETE_WINDOW", iptal)

    pencere.update_idletasks()
    x = pencere.winfo_screenwidth() // 2 - pencere.winfo_width() // 2
    y = pencere.winfo_screenheight() // 2 - pencere.winfo_height() // 2
    pencere.geometry(f"+{x}+{y}")

    root.wait_window(pencere)
    return sonuc["tcknler"]


def ayni_ipc_idli_girdileri_tekillestir(girdiler: Iterable[str]) -> Tuple[List[str], int]:
    """Ayni IPC id ile verilen kullanici girdilerini ilk gorulene indirger."""
    tekil_girdiler = []
    gorulen_ipc_idleri = set()
    atlanan_sayi = 0

    for girdi in girdiler:
        girdi_bilgisi = girdi_bilgisi_ayikla(girdi)
        ipc_id = str(girdi_bilgisi.get("ipc_id", "")).strip().upper()
        if ipc_id:
            if ipc_id in gorulen_ipc_idleri:
                atlanan_sayi += 1
                continue
            gorulen_ipc_idleri.add(ipc_id)
        tekil_girdiler.append(girdi)

    return tekil_girdiler, atlanan_sayi


def kimlik_turu_belirle(kimlik_no: str) -> str:
    """TCKN mi VKN mi belirler."""
    if len(kimlik_no) == 11:
        return "TCKN"
    elif len(kimlik_no) == 10:
        return "VKN"
    return "Bilinmiyor"


def log_dosya_yolu(base_dir: Path) -> Path:
    """Log dosyasının yolunu döndürür."""
    return base_dir / LOG_DOSYA_ADI


def log_workbook_hazirla(log_yolu: Path) -> Workbook:
    """Log workbook'unu hazırlar veya mevcut olanı yükler."""
    if log_yolu.exists():
        return load_workbook(log_yolu)
    wb = Workbook()
    ws = wb.active
    ws.title = "log"
    return wb


def log_kaydi_ekle(wb: Workbook, log_yolu: Path, sayfa_adi: str, veri: Dict[str, str]) -> None:
    """Log kaydını ekler."""
    if sayfa_adi not in wb.sheetnames:
        ws = wb.create_sheet(sayfa_adi)
        ws.append(list(veri.keys()))
    else:
        ws = wb[sayfa_adi]
    ws.append(list(veri.values()))
    wb.save(log_yolu)


# =========================
# ANA AKIŞ
# =========================

def main() -> None:
    base_dir = uygulama_klasoru()
    resource_dir = kaynak_klasoru()

    root = tk.Tk()
    root.withdraw()

    try:
        # Vergi Dairesi türü seçimi
        belge_islem = vergi_dairesi_turu_sec(root)
        if belge_islem is None:
            return

        # TCKN/VKN girişi
        tcknler = tcknleri_al(root)
        if tcknler is None:
            return
        tcknler, ayni_ipc_id_atlanan_sayi = ayni_ipc_idli_girdileri_tekillestir(tcknler)

        # Excel dosyasını seç
        excel_yolu = dosya_yoksa_sec(
            root,
            base_dir / EXCEL_DOSYA_ADI,
            "Excel dosyasını seçiniz",
            [("Excel dosyası", "*.xlsx")],
        )
        if excel_yolu is None:
            return

        kalan_alacak_yolu = dosya_yoksa_sec(
            root,
            base_dir / KALAN_ALACAK_DOSYA_ADI,
            "Kalan Alacak Dosyaları Dökümü dosyasını seçiniz",
            [("Excel dosyası", "*.xlsx")],
        )
        if kalan_alacak_yolu is None:
            return

        # Şablon dosyasını seç
        if belge_islem == "VERGI_DAIRESI_SPK":
            sablon_adi = VERGI_DAIRESI_SPK_SABLON_ADI
        elif belge_islem == "VERGI_DAIRESI_VD":
            sablon_adi = VERGI_DAIRESI_VD_SABLON_ADI

        sablon_yolu = dosya_yoksa_sec(
            root,
            resource_dir / SABLON_KLASOR_ADI / sablon_adi,
            "Word şablon dosyasını seçiniz",
            [("Word dosyası", "*.docx")],
        )
        if sablon_yolu is None:
            return

        # Log dosyasını hazırla
        log_yolu = log_dosya_yolu(base_dir)
        log_wb = log_workbook_hazirla(log_yolu)

        olusan_dosyalar: List[Path] = []
        hatalar: List[str] = []
        kalan_alacak_uyarilari: List[str] = []
        kalan_alacak_indeksi = kalan_alacak_indeksi_olustur(kalan_alacak_yolu)
        hazir_veriler: List[Dict[str, Any]] = []

        # Vergi Dairesi verilerini hazırla
        for tckn in tcknler:
            try:
                veri = excelden_vergi_dairesi_kaydi_bul_girdi(excel_yolu, tckn)
                secilen_unvan = kalan_alacak_unvan_eslesmesini_veriye_uygula(root, kalan_alacak_indeksi, tckn, veri)
                if secilen_unvan:
                    kalan_alacak_uyarilari.append(f"{tckn}: Kalan Alacak ünvan eşleşmesi seçildi: {secilen_unvan}")
                kalan_tutar, kalan_mesaj = kalan_alacak_ipc_tutari_bul(kalan_alacak_indeksi, tckn, veri)
                kalan_alacak_uyarilari.append(kalan_mesaj)
                if kalan_tutar:
                    veri["pc_tutari"] = kalan_tutar
                    for parcali in veri.get("parcali_tahsilatlar", []):
                        parcali["pc_tutari"] = kalan_tutar
                ek_ipcler = odemeyi_aciklayan_ipcleri_bul(kalan_alacak_indeksi, tckn, veri)
                if ek_ipcler:
                    if not ek_ipcleri_kullaniciya_sor(root, veri, ek_ipcler):
                        ek_ipcler = []
                else:
                    ek_ipcler = muhtemel_ipcleri_kullaniciya_sec(root, veri, muhtemel_ipcleri_bul(kalan_alacak_indeksi, tckn, veri))

                if ek_ipcler:
                    tahsil_tarihi = tarih_metni_parse(veri.get("tarih", ""))
                    ana_beklenen = ipc_beklenen_odeme_tutari(
                        {
                            "ipc": ipc_tarih_sayi_anahtari(veri.get("karar_tarihi", ""), veri.get("karar_no", "")),
                            "pc_tutari_decimal": para_decimal(veri.get("pc_tutari", "")),
                        },
                        tahsil_tarihi,
                    )
                    if ana_beklenen is not None:
                        veri["odenen_tutar"] = decimal_para_metni(ana_beklenen)
                    veri["ek_ipc_kayitlari"] = [
                        {
                            "muhatap": veri["muhatap"],
                            "kimlik_no": veri["kimlik_no"],
                            "karar_tarihi": kayit["karar_tarihi"],
                            "karar_no": kayit["karar_no"],
                            "pc_tutari": kayit["pc_tutari"],
                            "odenen_tutar": decimal_para_metni(kayit["beklenen_odeme"]),
                            "tahsilat_tarihi_fis_no": veri["tahsilat_tarihi_fis_no"],
                        }
                        for kayit in ek_ipcler
                    ]
                hazir_veriler.append(veri)
            except Exception as exc:
                hatalar.append(f"{tckn}: {exc}")

        vergi_dairesi_gruplari: Dict[str, List[Dict[str, Any]]] = {}
        for veri in hazir_veriler:
            grup_anahtari = vergi_dairesi_hitap_formatla(veri["vergi_dairesi_adi"])
            vergi_dairesi_gruplari.setdefault(grup_anahtari, []).append(veri)

        # Aynı hitaba giden kayıtları tek yazıda birleştir.
        for grup_anahtari, grup_verileri in vergi_dairesi_gruplari.items():
            try:
                ana_veri = dict(grup_verileri[0])
                ana_veri["birlesik_kayitlar"] = grup_verileri[1:]
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                if len(grup_verileri) == 1:
                    dosya_adi = f"{temiz_dosya_adi(ana_veri['muhatap'])}_{ana_veri['kimlik_no']}_vergi_dairesi_{timestamp}.docx"
                else:
                    dosya_adi = f"{temiz_dosya_adi(grup_anahtari)}_{len(grup_verileri)}_kayit_vergi_dairesi_{timestamp}.docx"
                cikti_yolu = base_dir / CIKTI_KLASOR_ADI / dosya_adi
                vergi_dairesi_word_olustur(sablon_yolu, cikti_yolu, ana_veri)
                olusan_dosyalar.append(cikti_yolu)

                for veri in grup_verileri:
                    log_kaydi_ekle(
                        log_wb,
                        log_yolu,
                        "vergi_dairesi",
                        {
                            "log_zamani": datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                            "yazi_turu": "Vergi Dairesi Yazısı",
                            "kimlik_turu": kimlik_turu_belirle(veri["kimlik_no"]),
                            "kimlik_no": veri["kimlik_no"],
                            "muhatap": muhatap_gorunum_formatla(veri["muhatap"], veri["kimlik_no"]),
                            "vergi_dairesi_adi": veri["vergi_dairesi_adi"],
                            "excel_sayfa": veri["excel_sayfa"],
                            "excel_satir": veri["excel_satir"],
                            "sablon": sablon_adi,
                            "olusan_dosya": str(cikti_yolu),
                            "karar_tarihi": veri["karar_tarihi"],
                            "karar_no": veri["karar_no"],
                            "pc_tutari": veri["pc_tutari"],
                            "odenen_tutar": veri["odenen_tutar"],
                            "tahsilat_tarihi_fis_no": veri["tahsilat_tarihi_fis_no"],
                        },
                    )
            except Exception as exc:
                hatalar.append(f"{grup_anahtari}: {exc}")

        # Sonuçları göster
        if olusan_dosyalar:
            mesaj = f"{len(olusan_dosyalar)} adet yazı oluşturuldu."
            if ayni_ipc_id_atlanan_sayi:
                mesaj += f"\nAynı IPC id ile girilen {ayni_ipc_id_atlanan_sayi} kayıt tekrar olduğu için atlandı."
            if hatalar:
                mesaj += f"\n{len(hatalar)} adet hata oluştu."
            if kalan_alacak_uyarilari:
                mesaj += "\n\nKalan Alacak arama sonucu:\n" + "\n".join(kalan_alacak_uyarilari)
            messagebox.showinfo("Başarılı", mesaj, parent=root)
        else:
            messagebox.showwarning("Uyarı", "Hiç yazı oluşturulamadı.", parent=root)

        if hatalar:
            hata_mesaji = "Hatalar:\n" + "\n".join(hatalar)
            messagebox.showerror("Hatalar", hata_mesaji, parent=root)

        # Oluşan dosyaları aç
        if olusan_dosyalar:
            try:
                import os
                if os.name == "nt":
                    os.startfile(str(olusan_dosyalar[0]))
            except Exception:
                pass

    except Exception as exc:
        messagebox.showerror("Hata", str(exc), parent=root)
        raise
    finally:
        try:
            root.destroy()
        except Exception:
            pass


if __name__ == "__main__":
    main()
