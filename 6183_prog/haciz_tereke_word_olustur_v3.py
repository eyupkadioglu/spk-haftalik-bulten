# -*- coding: utf-8 -*-
"""
Haciz / tereke yazısı Word oluşturucu

Kullanım:
1) Bu dosyayı haciz_tereke_takip.xlsx dosyasının bulunduğu klasöre koyun.
2) Aynı klasörde Word şablonu da bulunsun: TEK KİŞİVD.docx
3) Gerekirse şu paketleri kurun:
   pip install python-docx openpyxl
4) Çalıştırın:
   python haciz_tereke_word_olustur.py

Not:
- 6183 seçilirse Excel'in "6183" sayfasından veri alınır.
- Kolon eşleşmesi:
  B = Ad Soyad
  E = TCKN
  G = Hitap
  H = İlgi Tarihi
  I = İlgi Sayısı
- Word'deki sarı alanlar üretilen dosyada varsayılan olarak sarı bırakılmaz.
"""

from __future__ import annotations

import os
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:
    raise SystemExit("Eksik paket: python-docx. Kurulum: pip install python-docx")

try:
    from openpyxl import load_workbook
except ImportError:
    raise SystemExit("Eksik paket: openpyxl. Kurulum: pip install openpyxl")

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
except ImportError:
    raise SystemExit("tkinter bulunamadı. Windows Python kurulumunda tkinter genelde hazır gelir.")


# =========================
# AYARLAR
# =========================

EXCEL_DOSYA_ADI = "haciz_tereke_takip.xlsx"
SABLON_KLASOR_ADI = "sablonlar"
WORD_SABLON_ADI = "TEK KİŞİVD.docx"
MAHKEME_TEREKE_BAKIYELI_SABLON_ADI = "mahkemeterekecevapbakiyeli.docx"
MAHKEME_TEREKE_BAKIYESIZ_SABLON_ADI = "mahkemeterekecevapbakiyesiz.docx"
CIKTI_KLASOR_ADI = "olusan_word_yazilari"

# True: Word'deki sarı işaretler çıktı dosyasında temizlenir.
# False: Dinamik alanlar sarı vurgulu kalır.
FINALDE_SARI_ALANLARI_TEMIZLE = True


@dataclass
class BelgeAyari:
    secim_adi: str
    excel_sayfa_adi: Optional[str]
    aktif: bool
    aciklama: str = ""


BELGE_TURU_AYARLARI: Dict[str, BelgeAyari] = {
    "6183": BelgeAyari(
        secim_adi="6183 yazısı",
        excel_sayfa_adi="6183",
        aktif=True,
        aciklama="6183 yazısı",
    ),
    "MAHKEME_TEREKE": BelgeAyari(
        secim_adi="Mahkeme tereke yazısı",
        excel_sayfa_adi=None,
        aktif=True,
        aciklama="Mahkeme tereke yazısı",
    ),
}


# =========================
# GENEL YARDIMCI FONKSİYONLAR
# =========================

def uygulama_klasoru() -> Path:
    """Script hangi klasördeyse onu döndürür. PyInstaller ile paketlenirse exe klasörünü kullanır."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def temiz_dosya_adi(metin: str) -> str:
    metin = str(metin or "").strip()
    metin = re.sub(r"[\\/:*?\"<>|]+", "_", metin)
    metin = re.sub(r"\s+", "_", metin)
    return metin.strip("_") or "belge"


def normalize_tckn(value) -> str:
    """Excel/Tkinter girdisinden 11 haneli TCKN metni üretir."""
    if value is None:
        return ""

    if isinstance(value, float):
        if value.is_integer():
            value = int(value)

    text = str(value).strip()

    # Excel bazen 27884602366.0 gibi gösterebilir.
    if re.fullmatch(r"\d+\.0", text):
        text = text[:-2]

    digits = re.sub(r"\D", "", text)
    if not digits:
        return ""

    # TCKN normalde 11 hanelidir. Excel'de baştaki sıfır kaybolmuşsa tamamlar.
    return digits.zfill(11) if len(digits) <= 11 else digits


def hucre_metni(cell, tarih_mi: bool = False) -> str:
    """Openpyxl hücresini temiz metne çevirir."""
    value = cell.value
    if value is None:
        return ""

    if tarih_mi:
        if isinstance(value, (datetime, date)):
            return value.strftime("%d.%m.%Y")
        # Hücre tarih formatlı ise openpyxl genelde datetime döndürür; metinse olduğu gibi korunur.
        return str(value).strip()

    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value).strip()

    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")

    return str(value).strip()


def hitap_satirlarini_ayir(hitap: str) -> List[str]:
    """Excel G kolonundaki hitabı Word'deki satırlara böler."""
    hitap = str(hitap or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not hitap:
        return [""]
    satirlar = [s.strip() for s in hitap.split("\n") if s.strip()]
    return satirlar or [hitap]


def turkce_buyuk_harf(metin: str) -> str:
    """Türkçe i/ı dönüşümlerini koruyarak büyük harfe çevirir."""
    return str(metin or "").translate(str.maketrans({"i": "İ", "ı": "I"})).upper()


def paranteze_al(metin: str) -> str:
    metin = str(metin or "").strip()
    if not metin:
        return ""
    if metin.startswith("(") and metin.endswith(")"):
        return metin
    return f"({metin})"


def mahkeme_hitabi_olustur(mahkeme: str) -> str:
    hitap = turkce_buyuk_harf(mahkeme).strip()
    return re.sub(r"\s+HAKİMLİĞİNE\s*$", "", hitap).strip()


# =========================
# TKINTER ARAYÜZ
# =========================

def belge_turu_sec(root: tk.Tk) -> Optional[str]:
    """Kullanıcıya önce 6183 mü Mahkeme tereke mi sorar."""
    sonuc = {"secim": None}

    pencere = tk.Toplevel(root)
    pencere.title("Belge Türü Seçimi")
    pencere.resizable(False, False)
    pencere.grab_set()

    etiket = tk.Label(
        pencere,
        text="Oluşturulacak yazı türünü seçiniz:",
        font=("Segoe UI", 11),
        padx=24,
        pady=18,
    )
    etiket.pack()

    buton_frame = tk.Frame(pencere, padx=18, pady=12)
    buton_frame.pack(fill="x")

    def sec(deger: str):
        sonuc["secim"] = deger
        pencere.destroy()

    btn_6183 = tk.Button(
        buton_frame,
        text="6183 cevap yazısı",
        width=24,
        height=2,
        command=lambda: sec("6183"),
    )
    btn_6183.grid(row=0, column=0, padx=8, pady=6)

    btn_mahkeme = tk.Button(
        buton_frame,
        text="Mahkeme tereke cevap yazısı",
        width=24,
        height=2,
        command=lambda: sec("MAHKEME_TEREKE"),
    )
    btn_mahkeme.grid(row=0, column=1, padx=8, pady=6)

    def iptal():
        sonuc["secim"] = None
        pencere.destroy()

    pencere.protocol("WM_DELETE_WINDOW", iptal)

    pencere.update_idletasks()
    x = pencere.winfo_screenwidth() // 2 - pencere.winfo_width() // 2
    y = pencere.winfo_screenheight() // 2 - pencere.winfo_height() // 2
    pencere.geometry(f"+{x}+{y}")

    root.wait_window(pencere)
    return sonuc["secim"]


def mahkeme_tereke_alt_sec(root: tk.Tk) -> Optional[str]:
    sonuc = {"secim": None}

    pencere = tk.Toplevel(root)
    pencere.title("Mahkeme Tereke Seçimi")
    pencere.resizable(False, False)
    pencere.grab_set()

    etiket = tk.Label(
        pencere,
        text="Mahkeme tereke yazısı türünü seçiniz:",
        font=("Segoe UI", 11),
        padx=24,
        pady=18,
    )
    etiket.pack()

    buton_frame = tk.Frame(pencere, padx=18, pady=12)
    buton_frame.pack(fill="x")

    def sec(deger: str):
        sonuc["secim"] = deger
        pencere.destroy()

    tk.Button(
        buton_frame,
        text="Bakiyesi Var",
        width=24,
        height=2,
        command=lambda: sec("MAHKEME_TEREKE_BAKIYELI"),
    ).grid(row=0, column=0, padx=8, pady=6)

    tk.Button(
        buton_frame,
        text="Bakiyesi yok",
        width=24,
        height=2,
        command=lambda: sec("MAHKEME_TEREKE_BAKIYESIZ"),
    ).grid(row=0, column=1, padx=8, pady=6)

    def iptal():
        sonuc["secim"] = None
        pencere.destroy()

    pencere.protocol("WM_DELETE_WINDOW", iptal)
    pencere.update_idletasks()
    x = pencere.winfo_screenwidth() // 2 - pencere.winfo_width() // 2
    y = pencere.winfo_screenheight() // 2 - pencere.winfo_height() // 2
    pencere.geometry(f"+{x}+{y}")

    root.wait_window(pencere)
    return sonuc["secim"]


def tcknleri_ayir(metin: str) -> Tuple[List[str], List[str]]:
    """Virgül, boşluk veya satır sonuyla ayrılmış TCKN listesini temizler."""
    parcalar = [p for p in re.split(r"[\s,;]+", str(metin or "").strip()) if p]
    tcknler: List[str] = []
    hatali: List[str] = []
    gorulen = set()

    for parca in parcalar:
        tckn = normalize_tckn(parca)
        if len(tckn) != 11:
            hatali.append(parca)
            continue
        if tckn not in gorulen:
            tcknler.append(tckn)
            gorulen.add(tckn)

    return tcknler, hatali


def tcknleri_al(root: tk.Tk) -> Optional[List[str]]:
    while True:
        sonuc = {"metin": None}

        pencere = tk.Toplevel(root)
        pencere.title("TCKN Girişi")
        pencere.resizable(True, True)
        pencere.grab_set()

        etiket = tk.Label(
            pencere,
            text="TCKN'leri virgül, boşluk veya alt alta yazarak giriniz:",
            font=("Segoe UI", 10),
            padx=14,
            pady=10,
            anchor="w",
        )
        etiket.pack(fill="x")

        metin_alani = tk.Text(pencere, width=48, height=10, font=("Consolas", 10))
        metin_alani.pack(fill="both", expand=True, padx=14, pady=10)
        metin_alani.focus_set()

        buton_frame = tk.Frame(pencere, padx=14)
        buton_frame.pack(fill="x", pady=12)

        def tamam():
            sonuc["metin"] = metin_alani.get("1.0", "end")
            pencere.destroy()

        def iptal():
            sonuc["metin"] = None
            pencere.destroy()

        tk.Button(buton_frame, text="Tamam", width=12, command=tamam).pack(side="right", padx=6)
        tk.Button(buton_frame, text="İptal", width=12, command=iptal).pack(side="right")

        pencere.protocol("WM_DELETE_WINDOW", iptal)
        pencere.update_idletasks()
        x = pencere.winfo_screenwidth() // 2 - pencere.winfo_width() // 2
        y = pencere.winfo_screenheight() // 2 - pencere.winfo_height() // 2
        pencere.geometry(f"+{x}+{y}")

        root.wait_window(pencere)

        if sonuc["metin"] is None:
            return None

        tcknler, hatali = tcknleri_ayir(sonuc["metin"])
        if tcknler and not hatali:
            return tcknler

        mesajlar = []
        if not tcknler:
            mesajlar.append("En az bir geçerli 11 haneli TCKN giriniz.")
        if hatali:
            mesajlar.append("Hatalı girişler: " + ", ".join(hatali))
        messagebox.showwarning("Hatalı TCKN", "\n".join(mesajlar), parent=root)


def dosya_yoksa_sec(root: tk.Tk, beklenen_yol: Path, baslik: str, filetypes) -> Optional[Path]:
    if beklenen_yol.exists():
        return beklenen_yol

    messagebox.showwarning(
        "Dosya bulunamadı",
        f"Beklenen dosya bulunamadı:\n{beklenen_yol}\n\nLütfen dosyayı seçiniz.",
        parent=root,
    )
    secilen = filedialog.askopenfilename(title=baslik, filetypes=filetypes, parent=root)
    if not secilen:
        return None
    return Path(secilen)


# =========================
# EXCEL OKUMA
# =========================

def excelden_6183_kaydi_bul(excel_yolu: Path, tckn: str) -> Dict[str, str]:
    wb = load_workbook(excel_yolu, data_only=True)

    if "6183" not in wb.sheetnames:
        raise ValueError(f"Excel dosyasında '6183' adlı sayfa bulunamadı. Mevcut sayfalar: {', '.join(wb.sheetnames)}")

    ws = wb["6183"]

    for row in range(2, ws.max_row + 1):
        hucre_tckn = normalize_tckn(ws.cell(row=row, column=5).value)  # E kolonu
        if hucre_tckn == tckn:
            ad_soyad = hucre_metni(ws.cell(row=row, column=2))          # B kolonu
            hitap = hucre_metni(ws.cell(row=row, column=7))             # G kolonu
            ilgi_tarihi = hucre_metni(ws.cell(row=row, column=8), tarih_mi=True)  # H kolonu
            ilgi_sayisi = hucre_metni(ws.cell(row=row, column=9))       # I kolonu

            eksikler = []
            if not ad_soyad:
                eksikler.append("B / Ad Soyad")
            if not hitap:
                eksikler.append("G / Hitap")
            if not ilgi_tarihi:
                eksikler.append("H / İlgi Tarihi")
            if not ilgi_sayisi:
                eksikler.append("I / İlgi Sayısı")
            if eksikler:
                raise ValueError(
                    f"TCKN bulundu fakat {row}. satırda eksik veri var: " + ", ".join(eksikler)
                )

            return {
                "ad_soyad": ad_soyad,
                "tckn": hucre_tckn,
                "hitap": hitap,
                "ilgi_tarihi": ilgi_tarihi,
                "ilgi_sayisi": ilgi_sayisi,
                "excel_satir": str(row),
            }

    raise LookupError(f"TCKN bulunamadı: {tckn}")


def excelden_tereke_kaydi_bul(excel_yolu: Path, tckn: str) -> Dict[str, str]:
    wb = load_workbook(excel_yolu, data_only=True)

    if "tereke" not in wb.sheetnames:
        raise ValueError(f"Excel dosyasında 'tereke' adlı sayfa bulunamadı. Mevcut sayfalar: {', '.join(wb.sheetnames)}")

    ws = wb["tereke"]

    for row in range(2, ws.max_row + 1):
        hucre_tckn = normalize_tckn(ws.cell(row=row, column=3).value)  # C kolonu
        if hucre_tckn == tckn:
            muris_ad_soyad = hucre_metni(ws.cell(row=row, column=2))  # B kolonu
            mahkeme = hucre_metni(ws.cell(row=row, column=5))         # E kolonu
            dosya_no = hucre_metni(ws.cell(row=row, column=6))        # F kolonu
            tur = hucre_metni(ws.cell(row=row, column=7))             # G kolonu

            eksikler = []
            if not muris_ad_soyad:
                eksikler.append("B / Murisin Adı Soyadı")
            if not mahkeme:
                eksikler.append("E / Mahkeme")
            if not dosya_no:
                eksikler.append("F / Dosya No")
            if eksikler:
                raise ValueError(
                    f"TCKN bulundu fakat {row}. satırda eksik veri var: " + ", ".join(eksikler)
                )

            return {
                "muris_ad_soyad": muris_ad_soyad,
                "tckn": hucre_tckn,
                "mahkeme": mahkeme,
                "dosya_no": dosya_no,
                "tur": tur,
                "excel_satir": str(row),
            }

    raise LookupError(f"TCKN bulunamadı: {tckn}")


# =========================
# WORD BİÇİM YARDIMCILARI
# =========================

def run_format_snapshot(run) -> Dict[str, object]:
    """Bir run'ın temel biçimini saklar."""
    fmt = {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "style": run.style,
        "font_name": run.font.name,
        "font_size": run.font.size,
        "highlight": run.font.highlight_color,
        "color_rgb": None,
    }
    try:
        fmt["color_rgb"] = run.font.color.rgb
    except Exception:
        fmt["color_rgb"] = None
    return fmt


def paragraf_ilk_run_format(paragraph) -> Dict[str, object]:
    for run in paragraph.runs:
        if run.text:
            return run_format_snapshot(run)
    if paragraph.runs:
        return run_format_snapshot(paragraph.runs[0])
    # Güvenli varsayılan
    temp = paragraph.add_run("")
    fmt = run_format_snapshot(temp)
    paragraph._p.remove(temp._r)
    return fmt


def format_uygula(run, fmt: Dict[str, object], sariyi_temizle: bool = True) -> None:
    run.bold = fmt.get("bold")
    run.italic = fmt.get("italic")
    run.underline = fmt.get("underline")
    try:
        run.style = fmt.get("style")
    except Exception:
        pass
    run.font.name = fmt.get("font_name")
    run.font.size = fmt.get("font_size")
    if fmt.get("color_rgb") is not None:
        run.font.color.rgb = fmt.get("color_rgb")
    if sariyi_temizle:
        run.font.highlight_color = None
    else:
        run.font.highlight_color = fmt.get("highlight")


def paragrafi_temizle(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def paragraf_yaz(paragraph, text: str, fmt: Optional[Dict[str, object]] = None, sariyi_temizle: bool = True) -> None:
    """Paragrafın metnini değiştirir, paragraf stilini/yerleşimini korur."""
    if fmt is None:
        fmt = paragraf_ilk_run_format(paragraph)

    paragrafi_temizle(paragraph)

    satirlar = str(text or "").split("\n")
    if not satirlar:
        satirlar = [""]

    for i, satir in enumerate(satirlar):
        if i > 0:
            br = paragraph.add_run()
            format_uygula(br, fmt, sariyi_temizle=sariyi_temizle)
            br.add_break()
        run = paragraph.add_run(satir)
        format_uygula(run, fmt, sariyi_temizle=sariyi_temizle)


def paragrafta_sari_var_mi(paragraph) -> bool:
    return any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in paragraph.runs)


def hucre_ilk_paragraf(cell):
    if not cell.paragraphs:
        return cell.add_paragraph()
    return cell.paragraphs[0]


def sari_paragraflari_bul(doc: Document) -> List:
    paragraflar = []

    for p in doc.paragraphs:
        if paragrafta_sari_var_mi(p):
            paragraflar.append(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if paragrafta_sari_var_mi(p):
                        paragraflar.append(p)

    return paragraflar


def dosya_tur_metni(veri: Dict[str, str]) -> str:
    parcalar = [veri.get("dosya_no", "").strip(), veri.get("tur", "").strip()]
    return " ".join(parca for parca in parcalar if parca)


# =========================
# WORD ALANLARINI DOLDURMA
# =========================

def hitap_alani_doldur(doc: Document, hitap: str) -> None:
    satirlar = hitap_satirlarini_ayir(hitap)

    # Şablonda İlgi paragrafından önceki sarı paragraflar hitap alanlarıdır.
    hitap_paragraflari = []
    for p in doc.paragraphs:
        if p.text.strip().startswith("İlgi:"):
            break
        if p.text.strip() and paragrafta_sari_var_mi(p):
            hitap_paragraflari.append(p)

    if not hitap_paragraflari:
        raise ValueError("Word şablonunda hitap için sarı alan bulunamadı.")

    # 1. satır: örn. ADANA DEFTERDARLIĞI
    fmt1 = paragraf_ilk_run_format(hitap_paragraflari[0])
    paragraf_yaz(
        hitap_paragraflari[0],
        turkce_buyuk_harf(satirlar[0]),
        fmt=fmt1,
        sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE,
    )

    # 2. satır ve varsa devamı: örn. (Ziyapaşa Vergi Dairesi Müdürlüğü)
    if len(hitap_paragraflari) >= 2:
        ikinci_metin = "\n".join(paranteze_al(satir) for satir in satirlar[1:]) if len(satirlar) > 1 else ""
        fmt2 = paragraf_ilk_run_format(hitap_paragraflari[1])
        paragraf_yaz(
            hitap_paragraflari[1],
            ikinci_metin,
            fmt=fmt2,
            sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE,
        )
    elif len(satirlar) > 1:
        # Şablonda tek hitap paragrafı varsa kalan satırları aynı paragrafa alt satır olarak ekler.
        paragraf_yaz(
            hitap_paragraflari[0],
            "\n".join([turkce_buyuk_harf(satirlar[0]), *[paranteze_al(satir) for satir in satirlar[1:]]]),
            fmt=fmt1,
            sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE,
        )


def konu_alani_doldur(doc: Document, ad_soyad: str) -> None:
    yeni_konu = f"{ad_soyad} Hk."

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(text == "Konu" for text in row_text):
                hedef_cell = None
                for cell in row.cells:
                    if "Hk" in cell.text or paragrafta_sari_var_mi(hucre_ilk_paragraf(cell)):
                        hedef_cell = cell
                        break
                if hedef_cell is None and len(row.cells) >= 3:
                    hedef_cell = row.cells[2]
                if hedef_cell is None:
                    raise ValueError("Konu satırı bulundu fakat konu hücresi belirlenemedi.")

                p = hucre_ilk_paragraf(hedef_cell)
                fmt = paragraf_ilk_run_format(p)
                paragraf_yaz(p, yeni_konu, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
                return

    raise ValueError("Word şablonunda 'Konu' satırı bulunamadı.")


def ilgi_alani_doldur(doc: Document, ilgi_tarihi: str, ilgi_sayisi: str) -> None:
    yeni_ilgi = f"İlgi: {ilgi_tarihi} tarihli ve {ilgi_sayisi} sayılı yazınız."
    for p in doc.paragraphs:
        if p.text.strip().startswith("İlgi:"):
            fmt = paragraf_ilk_run_format(p)
            paragraf_yaz(p, yeni_ilgi, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
            return
    raise ValueError("Word şablonunda 'İlgi:' paragrafı bulunamadı.")


def ana_metin_alani_doldur(doc: Document, ad_soyad: str, tckn: str) -> None:
    yeni_metin = (
        f"İlgide kayıtlı yazınızda özetle, {ad_soyad} (TCKN:{tckn}) adlı yatırımcının "
        "vergi borçlarından dolayı Merkezimiz nezdinde yer alan hak ve alacaklarına "
        "6183 sayılı Amme Alacaklarının Tahsil Usulü Hakkında Kanun'un 79'uncu maddesi "
        "gereğince haciz konulduğu bildirilmiştir. "
    )

    for p in doc.paragraphs:
        if p.text.strip().startswith("İlgide kayıtlı yazınızda özetle"):
            fmt = paragraf_ilk_run_format(p)
            paragraf_yaz(p, yeni_metin, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
            return
    raise ValueError("Word şablonunda ana metin paragrafı bulunamadı.")


def tum_sari_vurgulari_temizle(doc: Document) -> None:
    """Kalan sarı vurguları temizler. Dinamik alanlar dışında sarı not bırakılmadıysa faydalıdır."""
    def temizle_paragraphs(paragraphs: Iterable):
        for p in paragraphs:
            for run in p.runs:
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    run.font.highlight_color = None

    temizle_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                temizle_paragraphs(cell.paragraphs)

    for section in doc.sections:
        temizle_paragraphs(section.header.paragraphs)
        temizle_paragraphs(section.footer.paragraphs)
        for table in section.header.tables + section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    temizle_paragraphs(cell.paragraphs)


def word_olustur(sablon_yolu: Path, cikti_yolu: Path, veri: Dict[str, str]) -> None:
    doc = Document(sablon_yolu)

    hitap_alani_doldur(doc, veri["hitap"])
    konu_alani_doldur(doc, veri["ad_soyad"])
    ilgi_alani_doldur(doc, veri["ilgi_tarihi"], veri["ilgi_sayisi"])
    ana_metin_alani_doldur(doc, veri["ad_soyad"], veri["tckn"])

    if FINALDE_SARI_ALANLARI_TEMIZLE:
        tum_sari_vurgulari_temizle(doc)

    cikti_yolu.parent.mkdir(parents=True, exist_ok=True)
    doc.save(cikti_yolu)


def mahkeme_tereke_word_olustur(
    sablon_yolu: Path,
    cikti_yolu: Path,
    veri: Dict[str, str],
    bakiyeli_mi: bool,
) -> None:
    doc = Document(sablon_yolu)

    muris = veri["muris_ad_soyad"]
    tckn = veri["tckn"]
    dosya_tur = dosya_tur_metni(veri)
    bugun = date.today().strftime("%d.%m.%Y")

    hitap = mahkeme_hitabi_olustur(veri["mahkeme"])
    ilgi = f"İlgi: {bugun} tarihli ve {dosya_tur} sayılı yazınız."
    konu = f"{dosya_tur} ({muris}) Hk."

    if bakiyeli_mi:
        alan_degerleri = [
            hitap,
            ilgi,
            (
                f"İlgide kayıtlı yazınızda Mahkemenizde görülmekte dava kapsamında {muris} "
                f"(TCKN:{tckn}) adlı murisin Merkezimiz uhdesindeki hesabına ilişkin bilgi ve "
                "belgelerin gönderilmesi talep edilmiştir. "
            ),
            (
                "İlgi’de kayıtlı yazınızda yer alan TCKN bilgisi kullanılarak Merkezimiz kayıtları "
                "üzerinde yapılan inceleme neticesinde, SPK'nun 83’üncü maddesinin dördüncü fıkrası "
                f"kapsamında {muris} adlı murisin Merkezimiz uhdesinde yer alan hesabında bulunan sermaye "
                "piyasası araçları ile sermaye piyasası araçlarının yazımız tarihi itibariyle değerlerini "
                "gösteren bakiye raporu (Ek.1) ile raporda yer alan bilgilere ilişkin açıklamalara "
                "(Ek.2) ekte yer verilmiştir. "
            ),
            f"EK: 1-Bakiye Raporu ({muris})",
            konu,
        ]
    else:
        alan_degerleri = [
            hitap,
            ilgi,
            (
                f"İlgide kayıtlı yazınız ile {muris} (TCKN:{tckn}) adlı murisin terekesinin iflas "
                "hükümlerine göre tasfiyesine karar verildiği ve tasfiye işlemlerine Mahkemenizin "
                f"{dosya_tur} sayılı dosyasından başlandığı TMK 622, İİK 226,180,181,184 ve "
                "208.maddeleri gereğince tebliğ edilmiştir."
            ),
            f"EK: Bakiye Raporu ({muris})",
            konu,
        ]

    sari_paragraflar = sari_paragraflari_bul(doc)
    if len(sari_paragraflar) < len(alan_degerleri):
        raise ValueError(
            f"Word şablonunda yeterli sarı alan bulunamadı. Beklenen: {len(alan_degerleri)}, bulunan: {len(sari_paragraflar)}"
        )

    for p, deger in zip(sari_paragraflar, alan_degerleri):
        fmt = paragraf_ilk_run_format(p)
        paragraf_yaz(p, deger, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)

    if FINALDE_SARI_ALANLARI_TEMIZLE:
        tum_sari_vurgulari_temizle(doc)

    cikti_yolu.parent.mkdir(parents=True, exist_ok=True)
    doc.save(cikti_yolu)


# =========================
# ANA AKIŞ
# =========================

def main() -> None:
    base_dir = uygulama_klasoru()

    root = tk.Tk()
    root.withdraw()

    try:
        secim = belge_turu_sec(root)
        if secim is None:
            return

        ayar = BELGE_TURU_AYARLARI[secim]
        if not ayar.aktif:
            messagebox.showinfo("Henüz tanımlı değil", ayar.aciklama, parent=root)
            return

        belge_islem = secim
        if secim == "MAHKEME_TEREKE":
            alt_secim = mahkeme_tereke_alt_sec(root)
            if alt_secim is None:
                return
            belge_islem = alt_secim

        tcknler = tcknleri_al(root)
        if tcknler is None:
            return

        excel_yolu = dosya_yoksa_sec(
            root,
            base_dir / EXCEL_DOSYA_ADI,
            "Excel dosyasını seçiniz",
            [("Excel dosyası", "*.xlsx")],
        )
        if excel_yolu is None:
            return

        if belge_islem == "MAHKEME_TEREKE_BAKIYELI":
            sablon_adi = MAHKEME_TEREKE_BAKIYELI_SABLON_ADI
        elif belge_islem == "MAHKEME_TEREKE_BAKIYESIZ":
            sablon_adi = MAHKEME_TEREKE_BAKIYESIZ_SABLON_ADI
        else:
            sablon_adi = WORD_SABLON_ADI

        sablon_yolu = dosya_yoksa_sec(
            root,
            base_dir / SABLON_KLASOR_ADI / sablon_adi,
            "Word şablon dosyasını seçiniz",
            [("Word dosyası", "*.docx")],
        )
        if sablon_yolu is None:
            return

        olusan_dosyalar: List[Path] = []
        hatalar: List[str] = []

        for tckn in tcknler:
            try:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

                if belge_islem == "MAHKEME_TEREKE_BAKIYELI":
                    veri = excelden_tereke_kaydi_bul(excel_yolu, tckn)
                    dosya_adi = f"{temiz_dosya_adi(veri['muris_ad_soyad'])}_{veri['tckn']}_tereke_bakiyeli_{timestamp}.docx"
                    cikti_yolu = base_dir / CIKTI_KLASOR_ADI / dosya_adi
                    mahkeme_tereke_word_olustur(sablon_yolu, cikti_yolu, veri, bakiyeli_mi=True)
                elif belge_islem == "MAHKEME_TEREKE_BAKIYESIZ":
                    veri = excelden_tereke_kaydi_bul(excel_yolu, tckn)
                    dosya_adi = f"{temiz_dosya_adi(veri['muris_ad_soyad'])}_{veri['tckn']}_tereke_bakiyesiz_{timestamp}.docx"
                    cikti_yolu = base_dir / CIKTI_KLASOR_ADI / dosya_adi
                    mahkeme_tereke_word_olustur(sablon_yolu, cikti_yolu, veri, bakiyeli_mi=False)
                else:
                    veri = excelden_6183_kaydi_bul(excel_yolu, tckn)
                    dosya_adi = f"{temiz_dosya_adi(veri['ad_soyad'])}_{veri['tckn']}_6183_{timestamp}.docx"
                    cikti_yolu = base_dir / CIKTI_KLASOR_ADI / dosya_adi
                    word_olustur(sablon_yolu, cikti_yolu, veri)

                olusan_dosyalar.append(cikti_yolu)
            except Exception as exc:
                hatalar.append(f"{tckn}: {exc}")

        mesaj = [f"Oluşturulan Word dosyası sayısı: {len(olusan_dosyalar)}"]
        if olusan_dosyalar:
            mesaj.append("")
            mesaj.extend(str(yol) for yol in olusan_dosyalar)
        if hatalar:
            mesaj.append("")
            mesaj.append("Oluşturulamayan kayıtlar:")
            mesaj.extend(hatalar)

        messagebox.showinfo("İşlem tamamlandı", "\n".join(mesaj), parent=root)

        # Windows'ta tek dosya oluştuysa dosyayı, birden fazla dosya oluştuysa çıktı klasörünü açar.
        try:
            if os.name == "nt" and olusan_dosyalar:
                acilacak_yol = olusan_dosyalar[0] if len(olusan_dosyalar) == 1 else olusan_dosyalar[0].parent
                os.startfile(str(acilacak_yol))  # type: ignore[attr-defined]
        except Exception:
            pass

    except Exception as exc:
        messagebox.showerror("Hata", str(exc), parent=root)
        raise
    finally:
        try:
            root.destroy()
        except Exception:
            pass


if __name__ == "__main__":
    main()
