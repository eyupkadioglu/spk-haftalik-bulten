# -*- coding: utf-8 -*-
"""
Haciz / tereke yazısı Word oluşturucu
Copyright: Eyüp Kadıoğlu 2026

Kullanım:
1) Bu dosyayı haciz_tereke_takip.xlsx dosyasının bulunduğu klasöre koyun.
2) Aynı klasörde Word şablonu da bulunsun: TEK KİŞİVD.docx
3) Gerekirse şu paketleri kurun:
   pip install python-docx openpyxl
4) Çalıştırın:
   python haciz_tereke_word_olustur.py

Not:
- 6183 seçilirse Excel'in "6183" sayfasından veri alınır.
- Kolon eşleşmesi:
  B = Ad Soyad
  D = Gerçek / Tüzel
  E = TCKN / VKN
  F = Hitap
  G = İlgi Tarihi
  H = İlgi Sayısı
- Word'deki sarı alanlar üretilen dosyada varsayılan olarak sarı bırakılmaz.
"""

from __future__ import annotations

import os
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:
    raise SystemExit("Eksik paket: python-docx. Kurulum: pip install python-docx")

try:
    from openpyxl import Workbook, load_workbook
except ImportError:
    raise SystemExit("Eksik paket: openpyxl. Kurulum: pip install openpyxl")

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
except ImportError:
    raise SystemExit("tkinter bulunamadı. Windows Python kurulumunda tkinter genelde hazır gelir.")


# =========================
# AYARLAR
# =========================

EXCEL_DOSYA_ADI = "haciz_tereke_takip.xlsx"
SABLON_KLASOR_ADI = "sablonlar"
WORD_SABLON_ADI = "TEK KİŞİVD.docx"
MAHKEME_TEREKE_BAKIYELI_SABLON_ADI = "mahkemeterekecevapbakiyeli.docx"
MAHKEME_TEREKE_BAKIYESIZ_SABLON_ADI = "mahkemeterekecevapbakiyesiz.docx"
TAKASBANK_BILDIRIM_SABLON_ADI = "takasbankbildirim.docx"
YTM_ICI_TALEP_SABLON_ADI = "ytmicisorgutablo.docx"
CIKTI_KLASOR_ADI = "olusan_word_yazilari"
LOG_DOSYA_ADI = "log_kaydi.xlsx"

# True: Word'deki sarı işaretler çıktı dosyasında temizlenir.
# False: Dinamik alanlar sarı vurgulu kalır.
FINALDE_SARI_ALANLARI_TEMIZLE = True
COPYRIGHT = "Eyüp Kadıoğlu 2026"


@dataclass
class BelgeAyari:
    secim_adi: str
    excel_sayfa_adi: Optional[str]
    aktif: bool
    aciklama: str = ""


BELGE_TURU_AYARLARI: Dict[str, BelgeAyari] = {
    "6183": BelgeAyari(
        secim_adi="6183 yazısı",
        excel_sayfa_adi="6183",
        aktif=True,
        aciklama="6183 yazısı",
    ),
    "MAHKEME_TEREKE": BelgeAyari(
        secim_adi="Mahkeme tereke yazısı",
        excel_sayfa_adi=None,
        aktif=True,
        aciklama="Mahkeme tereke yazısı",
    ),
    "TAKASBANK_BILDIRIM": BelgeAyari(
        secim_adi="Takasbank Bildirim yazısı",
        excel_sayfa_adi="6183",
        aktif=True,
        aciklama="Takasbank Bildirim yazısı",
    ),
    "YTM_ICI_TALEP": BelgeAyari(
        secim_adi="YTM İçi Talep Yazısı",
        excel_sayfa_adi="tereke",
        aktif=True,
        aciklama="YTM İçi Talep Yazısı",
    ),
}


# =========================
# GENEL YARDIMCI FONKSİYONLAR
# =========================

def uygulama_klasoru() -> Path:
    """Script hangi klasördeyse onu döndürür. PyInstaller ile paketlenirse exe klasörünü kullanır."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def temiz_dosya_adi(metin: str) -> str:
    metin = str(metin or "").strip()
    metin = re.sub(r"[\\/:*?\"<>|]+", "_", metin)
    metin = re.sub(r"\s+", "_", metin)
    return metin.strip("_") or "belge"


def normalize_tckn(value) -> str:
    """Excel/Tkinter girdisinden 11 haneli TCKN metni üretir."""
    if value is None:
        return ""

    if isinstance(value, float):
        if value.is_integer():
            value = int(value)

    text = str(value).strip()

    # Excel bazen 27884602366.0 gibi gösterebilir.
    if re.fullmatch(r"\d+\.0", text):
        text = text[:-2]

    digits = re.sub(r"\D", "", text)
    if not digits:
        return ""

    # TCKN normalde 11 hanelidir. Excel'de baştaki sıfır kaybolmuşsa tamamlar.
    return digits.zfill(11) if len(digits) <= 11 else digits


def normalize_tckn_vkn(value) -> str:
    """Excel/Tkinter girdisinden 10 haneli VKN veya 11 haneli TCKN metni üretir."""
    if value is None:
        return ""

    if isinstance(value, float):
        if value.is_integer():
            value = int(value)

    text = str(value).strip()
    if re.fullmatch(r"\d+\.0", text):
        text = text[:-2]

    digits = re.sub(r"\D", "", text)
    if not digits:
        return ""
    if len(digits) < 10:
        return digits.zfill(10)
    return digits


def kullanici_kimlik_no_denetle(value: str) -> Tuple[Optional[str], Optional[str]]:
    text = str(value or "").strip()
    digits = re.sub(r"\D", "", text)

    if not digits:
        return None, "rakam içermiyor"
    if len(digits) == 10:
        return digits, None
    if len(digits) == 11:
        if int(digits[-1]) % 2 != 0:
            return None, "TCKN son hanesi çift sayı olmalı"
        return digits, None
    return None, "VKN 10 hane, TCKN 11 hane olmalı"


def hucre_metni(cell, tarih_mi: bool = False) -> str:
    """Openpyxl hücresini temiz metne çevirir."""
    value = cell.value
    if value is None:
        return ""

    if tarih_mi:
        if isinstance(value, (datetime, date)):
            return value.strftime("%d.%m.%Y")
        # Hücre tarih formatlı ise openpyxl genelde datetime döndürür; metinse olduğu gibi korunur.
        return str(value).strip()

    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value).strip()

    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")

    return str(value).strip()


def son_dolu_satir(ws, kolonlar: Iterable[int]) -> int:
    """Sadece verilen kolonlarda gerçek değer bulunan son satırı döndürür."""
    kolon_seti = set(kolonlar)
    dolu_satirlar = [
        row
        for (row, col), cell in ws._cells.items()
        if col in kolon_seti and cell.value not in (None, "")
    ]
    return max(dolu_satirlar, default=1)


def dosya_no_tur_ayir(metin: str) -> Tuple[str, str]:
    """F kolonundaki birleşik 'dosya no + tür' değerini log için ayırır."""
    metin = str(metin or "").strip()
    if not metin:
        return "", ""
    parcalar = metin.split(maxsplit=1)
    if len(parcalar) == 1:
        return parcalar[0], ""
    return parcalar[0], parcalar[1]


def hitap_satirlarini_ayir(hitap: str) -> List[str]:
    """Excel G kolonundaki hitabı Word'deki satırlara böler."""
    hitap = str(hitap or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not hitap:
        return [""]
    satirlar = [s.strip() for s in hitap.split("\n") if s.strip()]
    return satirlar or [hitap]


def turkce_buyuk_harf(metin: str) -> str:
    """Türkçe i/ı dönüşümlerini koruyarak büyük harfe çevirir."""
    return str(metin or "").translate(str.maketrans({"i": "İ", "ı": "I"})).upper()


def paranteze_al(metin: str) -> str:
    metin = str(metin or "").strip()
    if not metin:
        return ""
    if metin.startswith("(") and metin.endswith(")"):
        return metin
    return f"({metin})"


def mahkeme_hitabi_olustur(mahkeme: str) -> str:
    hitap = turkce_buyuk_harf(mahkeme).strip()
    return re.sub(r"\s+HAKİMLİĞİNE\s*$", "", hitap).strip()


# =========================
# TKINTER ARAYÜZ
# =========================

def copyright_etiketi_ekle(pencere) -> None:
    tk.Label(
        pencere,
        text=COPYRIGHT,
        font=("Segoe UI", 8),
        fg="#666666",
        padx=12,
        pady=6,
    ).pack(side="bottom", fill="x")


def belge_turu_sec(root: tk.Tk) -> Optional[str]:
    """Kullanıcıya önce 6183 mü Mahkeme tereke mi sorar."""
    sonuc = {"secim": None}

    pencere = tk.Toplevel(root)
    pencere.title("Belge Türü Seçimi")
    pencere.resizable(False, False)
    pencere.grab_set()

    etiket = tk.Label(
        pencere,
        text="Oluşturulacak yazı türünü seçiniz:",
        font=("Segoe UI", 11),
        padx=24,
        pady=18,
    )
    etiket.pack()

    buton_frame = tk.Frame(pencere, padx=18, pady=12)
    buton_frame.pack(fill="x")

    def sec(deger: str):
        sonuc["secim"] = deger
        pencere.destroy()

    btn_6183 = tk.Button(
        buton_frame,
        text="6183 cevap yazısı",
        width=24,
        height=2,
        command=lambda: sec("6183"),
    )
    btn_6183.grid(row=0, column=0, padx=8, pady=6)

    btn_mahkeme = tk.Button(
        buton_frame,
        text="Mahkeme tereke cevap yazısı",
        width=24,
        height=2,
        command=lambda: sec("MAHKEME_TEREKE"),
    )
    btn_mahkeme.grid(row=0, column=1, padx=8, pady=6)

    btn_takasbank = tk.Button(
        buton_frame,
        text="Takasbank Bildirim yazısı",
        width=24,
        height=2,
        command=lambda: sec("TAKASBANK_BILDIRIM"),
    )
    btn_takasbank.grid(row=1, column=0, padx=8, pady=6)

    btn_ytm_ici = tk.Button(
        buton_frame,
        text="YTM İçi Talep Yazısı",
        width=24,
        height=2,
        command=lambda: sec("YTM_ICI_TALEP"),
    )
    btn_ytm_ici.grid(row=1, column=1, padx=8, pady=6)
    copyright_etiketi_ekle(pencere)

    def iptal():
        sonuc["secim"] = None
        pencere.destroy()

    pencere.protocol("WM_DELETE_WINDOW", iptal)

    pencere.update_idletasks()
    x = pencere.winfo_screenwidth() // 2 - pencere.winfo_width() // 2
    y = pencere.winfo_screenheight() // 2 - pencere.winfo_height() // 2
    pencere.geometry(f"+{x}+{y}")

    root.wait_window(pencere)
    return sonuc["secim"]


def mahkeme_tereke_alt_sec(root: tk.Tk) -> Optional[str]:
    sonuc = {"secim": None}

    pencere = tk.Toplevel(root)
    pencere.title("Mahkeme Tereke Seçimi")
    pencere.resizable(False, False)
    pencere.grab_set()

    etiket = tk.Label(
        pencere,
        text="Mahkeme tereke yazısı türünü seçiniz:",
        font=("Segoe UI", 11),
        padx=24,
        pady=18,
    )
    etiket.pack()

    buton_frame = tk.Frame(pencere, padx=18, pady=12)
    buton_frame.pack(fill="x")

    def sec(deger: str):
        sonuc["secim"] = deger
        pencere.destroy()

    tk.Button(
        buton_frame,
        text="Bakiyesi Var",
        width=24,
        height=2,
        command=lambda: sec("MAHKEME_TEREKE_BAKIYELI"),
    ).grid(row=0, column=0, padx=8, pady=6)

    tk.Button(
        buton_frame,
        text="Bakiyesi yok",
        width=24,
        height=2,
        command=lambda: sec("MAHKEME_TEREKE_BAKIYESIZ"),
    ).grid(row=0, column=1, padx=8, pady=6)
    copyright_etiketi_ekle(pencere)

    def iptal():
        sonuc["secim"] = None
        pencere.destroy()

    pencere.protocol("WM_DELETE_WINDOW", iptal)
    pencere.update_idletasks()
    x = pencere.winfo_screenwidth() // 2 - pencere.winfo_width() // 2
    y = pencere.winfo_screenheight() // 2 - pencere.winfo_height() // 2
    pencere.geometry(f"+{x}+{y}")

    root.wait_window(pencere)
    return sonuc["secim"]


def tcknleri_ayir(metin: str) -> Tuple[List[str], List[str]]:
    """Virgül, boşluk veya satır sonuyla ayrılmış TCKN/VKN listesini temizler."""
    parcalar = [p for p in re.split(r"[\s,;]+", str(metin or "").strip()) if p]
    tcknler: List[str] = []
    hatali: List[str] = []
    gorulen = set()

    for parca in parcalar:
        tckn, hata = kullanici_kimlik_no_denetle(parca)
        if hata:
            hatali.append(f"{parca} ({hata})")
            continue
        if tckn not in gorulen:
            tcknler.append(tckn)
            gorulen.add(tckn)

    return tcknler, hatali


def tcknleri_al(root: tk.Tk) -> Optional[List[str]]:
    while True:
        sonuc = {"metin": None}

        pencere = tk.Toplevel(root)
        pencere.title("TCKN Girişi")
        pencere.resizable(True, True)
        pencere.grab_set()

        etiket = tk.Label(
            pencere,
            text="TCKN/VKN'leri virgül, boşluk veya alt alta yazarak giriniz:",
            font=("Segoe UI", 10),
            padx=14,
            pady=10,
            anchor="w",
        )
        etiket.pack(fill="x")

        metin_alani = tk.Text(pencere, width=48, height=10, font=("Consolas", 10))
        metin_alani.pack(fill="both", expand=True, padx=14, pady=10)
        metin_alani.focus_set()

        buton_frame = tk.Frame(pencere, padx=14)
        buton_frame.pack(fill="x", pady=12)

        def tamam():
            sonuc["metin"] = metin_alani.get("1.0", "end")
            pencere.destroy()

        def iptal():
            sonuc["metin"] = None
            pencere.destroy()

        tk.Button(buton_frame, text="Tamam", width=12, command=tamam).pack(side="right", padx=6)
        tk.Button(buton_frame, text="İptal", width=12, command=iptal).pack(side="right")
        copyright_etiketi_ekle(pencere)

        pencere.protocol("WM_DELETE_WINDOW", iptal)
        pencere.update_idletasks()
        x = pencere.winfo_screenwidth() // 2 - pencere.winfo_width() // 2
        y = pencere.winfo_screenheight() // 2 - pencere.winfo_height() // 2
        pencere.geometry(f"+{x}+{y}")

        root.wait_window(pencere)

        if sonuc["metin"] is None:
            return None

        tcknler, hatali = tcknleri_ayir(sonuc["metin"])
        if tcknler and not hatali:
            return tcknler

        mesajlar = []
        if not tcknler:
            mesajlar.append("En az bir geçerli 10 haneli VKN veya son hanesi çift olan 11 haneli TCKN giriniz.")
        if hatali:
            mesajlar.append("Hatalı girişler: " + ", ".join(hatali))
        messagebox.showwarning("Hatalı TCKN/VKN", "\n".join(mesajlar), parent=root)


def dosya_yoksa_sec(root: tk.Tk, beklenen_yol: Path, baslik: str, filetypes) -> Optional[Path]:
    if beklenen_yol.exists():
        return beklenen_yol

    messagebox.showwarning(
        "Dosya bulunamadı",
        f"Beklenen dosya bulunamadı:\n{beklenen_yol}\n\nLütfen dosyayı seçiniz.",
        parent=root,
    )
    secilen = filedialog.askopenfilename(title=baslik, filetypes=filetypes, parent=root)
    if not secilen:
        return None
    return Path(secilen)


# =========================
# EXCEL OKUMA
# =========================

def excelden_6183_kaydi_bul(excel_yolu: Path, kimlik_no: str) -> Dict[str, str]:
    wb = load_workbook(excel_yolu, data_only=True)

    if "6183" not in wb.sheetnames:
        raise ValueError(f"Excel dosyasında '6183' adlı sayfa bulunamadı. Mevcut sayfalar: {', '.join(wb.sheetnames)}")

    ws = wb["6183"]
    eslesen_satirlar = []

    for row in range(2, ws.max_row + 1):
        hucre_kimlik_no = normalize_tckn_vkn(ws.cell(row=row, column=5).value)  # E kolonu
        if hucre_kimlik_no == kimlik_no:
            eslesen_satirlar.append(row)

    if not eslesen_satirlar:
        raise LookupError(f"TCKN/VKN bulunamadı: {kimlik_no}")

    row = max(eslesen_satirlar)
    ad_soyad = hucre_metni(ws.cell(row=row, column=2))          # B kolonu
    kisi_turu = hucre_metni(ws.cell(row=row, column=4))         # D kolonu
    hitap = hucre_metni(ws.cell(row=row, column=6))             # F kolonu
    ilgi_tarihi = hucre_metni(ws.cell(row=row, column=7), tarih_mi=True)  # G kolonu
    ilgi_sayisi = hucre_metni(ws.cell(row=row, column=8))       # H kolonu

    eksikler = []
    if not ad_soyad:
        eksikler.append("B / Ad Soyad")
    if not kisi_turu:
        eksikler.append("D / Gerçek-Tüzel")
    if not hitap:
        eksikler.append("F / Hitap")
    if not ilgi_tarihi:
        eksikler.append("G / İlgi Tarihi")
    if not ilgi_sayisi:
        eksikler.append("H / İlgi Sayısı")
    if eksikler:
        raise ValueError(
            f"TCKN/VKN bulundu fakat {row}. satırda eksik veri var: " + ", ".join(eksikler)
        )

    onceki_satirlar = [str(satir) for satir in eslesen_satirlar if satir != row]
    return {
        "ad_soyad": ad_soyad,
        "kimlik_no": normalize_tckn_vkn(ws.cell(row=row, column=5).value),
        "kisi_turu": kisi_turu,
        "hitap": hitap,
        "ilgi_tarihi": ilgi_tarihi,
        "ilgi_sayisi": ilgi_sayisi,
        "excel_sayfa": "6183",
        "excel_satir": str(row),
        "onceki_excel_satirlari": ", ".join(onceki_satirlar),
    }


def excelden_tereke_kaydi_bul(excel_yolu: Path, kimlik_no: str) -> Dict[str, str]:
    wb = load_workbook(excel_yolu, data_only=True)

    if "tereke" not in wb.sheetnames:
        raise ValueError(f"Excel dosyasında 'tereke' adlı sayfa bulunamadı. Mevcut sayfalar: {', '.join(wb.sheetnames)}")

    ws = wb["tereke"]
    eslesen_satirlar = []
    son_satir = son_dolu_satir(ws, (2, 3, 4, 5, 6))

    for row in range(2, son_satir + 1):
        hucre_kimlik_no = normalize_tckn_vkn(ws.cell(row=row, column=3).value)  # C kolonu
        if hucre_kimlik_no == kimlik_no:
            eslesen_satirlar.append(row)

    if not eslesen_satirlar:
        raise LookupError(f"TCKN/VKN bulunamadı: {kimlik_no}")

    row = max(eslesen_satirlar)
    muris_ad_soyad = hucre_metni(ws.cell(row=row, column=2))          # B kolonu
    mahkeme = hucre_metni(ws.cell(row=row, column=4))                 # D kolonu
    ilgi_tarihi = hucre_metni(ws.cell(row=row, column=5), tarih_mi=True)  # E kolonu
    dosya_tur = hucre_metni(ws.cell(row=row, column=6))               # F kolonu
    dosya_no, tur = dosya_no_tur_ayir(dosya_tur)

    eksikler = []
    if not muris_ad_soyad:
        eksikler.append("B / Murisin Adı Soyadı")
    if not mahkeme:
        eksikler.append("D / Mahkeme")
    if not ilgi_tarihi:
        eksikler.append("E / Tarih")
    if not dosya_tur:
        eksikler.append("F / Dosya No ve Tür")
    if eksikler:
        raise ValueError(
            f"TCKN/VKN bulundu fakat {row}. satırda eksik veri var: " + ", ".join(eksikler)
        )

    onceki_satirlar = [str(satir) for satir in eslesen_satirlar if satir != row]
    normalize_kimlik_no = normalize_tckn_vkn(ws.cell(row=row, column=3).value)
    return {
        "muris_ad_soyad": muris_ad_soyad,
        "tckn": normalize_kimlik_no,
        "kimlik_no": normalize_kimlik_no,
        "kimlik_turu": kimlik_turu_belirle(normalize_kimlik_no),
        "mahkeme": mahkeme,
        "ilgi_tarihi": ilgi_tarihi,
        "dosya_tur": dosya_tur,
        "dosya_no": dosya_no,
        "tur": tur,
        "excel_sayfa": "tereke",
        "excel_satir": str(row),
        "onceki_excel_satirlari": ", ".join(onceki_satirlar),
    }


# =========================
# WORD BİÇİM YARDIMCILARI
# =========================

def run_format_snapshot(run) -> Dict[str, object]:
    """Bir run'ın temel biçimini saklar."""
    fmt = {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "style": run.style,
        "font_name": run.font.name,
        "font_size": run.font.size,
        "highlight": run.font.highlight_color,
        "color_rgb": None,
    }
    try:
        fmt["color_rgb"] = run.font.color.rgb
    except Exception:
        fmt["color_rgb"] = None
    return fmt


def paragraf_ilk_run_format(paragraph) -> Dict[str, object]:
    for run in paragraph.runs:
        if run.text:
            return run_format_snapshot(run)
    if paragraph.runs:
        return run_format_snapshot(paragraph.runs[0])
    # Güvenli varsayılan
    temp = paragraph.add_run("")
    fmt = run_format_snapshot(temp)
    paragraph._p.remove(temp._r)
    return fmt


def format_uygula(run, fmt: Dict[str, object], sariyi_temizle: bool = True) -> None:
    run.bold = fmt.get("bold")
    run.italic = fmt.get("italic")
    run.underline = fmt.get("underline")
    try:
        run.style = fmt.get("style")
    except Exception:
        pass
    run.font.name = fmt.get("font_name")
    run.font.size = fmt.get("font_size")
    if fmt.get("color_rgb") is not None:
        run.font.color.rgb = fmt.get("color_rgb")
    if sariyi_temizle:
        run.font.highlight_color = None
    else:
        run.font.highlight_color = fmt.get("highlight")


def paragrafi_temizle(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def paragraf_yaz(paragraph, text: str, fmt: Optional[Dict[str, object]] = None, sariyi_temizle: bool = True) -> None:
    """Paragrafın metnini değiştirir, paragraf stilini/yerleşimini korur."""
    if fmt is None:
        fmt = paragraf_ilk_run_format(paragraph)

    paragrafi_temizle(paragraph)

    satirlar = str(text or "").split("\n")
    if not satirlar:
        satirlar = [""]

    for i, satir in enumerate(satirlar):
        if i > 0:
            br = paragraph.add_run()
            format_uygula(br, fmt, sariyi_temizle=sariyi_temizle)
            br.add_break()
        run = paragraph.add_run(satir)
        format_uygula(run, fmt, sariyi_temizle=sariyi_temizle)


def paragrafta_sari_var_mi(paragraph) -> bool:
    return any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in paragraph.runs)


def hucre_ilk_paragraf(cell):
    if not cell.paragraphs:
        return cell.add_paragraph()
    return cell.paragraphs[0]


def hucre_yaz(cell, text: str) -> None:
    p = hucre_ilk_paragraf(cell)
    fmt = paragraf_ilk_run_format(p)
    paragraf_yaz(p, text, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
    for fazla_p in list(cell.paragraphs[1:]):
        fazla_p._element.getparent().remove(fazla_p._element)


def sari_paragraflari_bul(doc: Document) -> List:
    paragraflar = []

    for p in doc.paragraphs:
        if paragrafta_sari_var_mi(p):
            paragraflar.append(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if paragrafta_sari_var_mi(p):
                        paragraflar.append(p)

    return paragraflar


def dosya_tur_metni(veri: Dict[str, str]) -> str:
    if veri.get("dosya_tur", "").strip():
        return veri["dosya_tur"].strip()
    parcalar = [veri.get("dosya_no", "").strip(), veri.get("tur", "").strip()]
    return " ".join(parca for parca in parcalar if parca)


# =========================
# LOG YARDIMCILARI
# =========================

LOG_SAYFA_BASLIKLARI: Dict[str, List[str]] = {
    "6183": [
        "log_zamani",
        "yazi_turu",
        "kimlik_turu",
        "kimlik_no",
        "kisi_turu",
        "ad_unvan",
        "excel_sayfa",
        "excel_satir",
        "sablon",
        "olusan_dosya",
        "hitap",
        "ilgi_tarihi",
        "ilgi_sayisi",
    ],
    "mahkeme_tereke": [
        "log_zamani",
        "yazi_turu",
        "alt_tur",
        "kimlik_turu",
        "kimlik_no",
        "muris_ad_soyad",
        "excel_sayfa",
        "excel_satir",
        "sablon",
        "olusan_dosya",
        "mahkeme",
        "dosya_no",
        "tur",
    ],
    "takasbank": [
        "log_zamani",
        "yazi_turu",
        "kimlik_turu",
        "kimlik_no",
        "kisi_turu",
        "ad_unvan",
        "excel_sayfa",
        "excel_satir",
        "sablon",
        "olusan_dosya",
        "grup_no",
        "haczi_koyan",
        "haciz_bildirim_tarihi",
        "haciz_bildirim_sayisi",
    ],
    "ytm_ici_talep": [
        "log_zamani",
        "yazi_turu",
        "kimlik_turu",
        "kimlik_no",
        "muris_ad_soyad",
        "excel_sayfa",
        "excel_satir",
        "sablon",
        "olusan_dosya",
        "sira_no",
        "mahkeme",
        "ilgi_tarihi",
        "dosya_no",
        "tur",
    ],
}


def kimlik_turu_belirle(kimlik_no: str) -> str:
    return "VKN" if len(str(kimlik_no)) == 10 else "TCKN"


def log_dosya_yolu(base_dir: Path) -> Path:
    return base_dir / SABLON_KLASOR_ADI / LOG_DOSYA_ADI


def log_workbook_hazirla(log_yolu: Path):
    log_yolu.parent.mkdir(parents=True, exist_ok=True)
    if log_yolu.exists():
        wb = load_workbook(log_yolu)
    else:
        wb = Workbook()
        if "Sheet" in wb.sheetnames:
            wb.remove(wb["Sheet"])

    for sayfa_adi, basliklar in LOG_SAYFA_BASLIKLARI.items():
        ws = wb[sayfa_adi] if sayfa_adi in wb.sheetnames else wb.create_sheet(sayfa_adi)
        if ws.max_row == 1 and all(ws.cell(row=1, column=col).value is None for col in range(1, len(basliklar) + 1)):
            ws.append(basliklar)
        else:
            mevcut = [ws.cell(row=1, column=col).value for col in range(1, len(basliklar) + 1)]
            if mevcut != basliklar:
                for col, baslik in enumerate(basliklar, start=1):
                    ws.cell(row=1, column=col).value = baslik

    wb.save(log_yolu)
    return wb


def log_kullanilan_satirlar(wb, sayfa_adi: str, kimlik_no: str) -> set:
    if sayfa_adi not in wb.sheetnames:
        return set()
    ws = wb[sayfa_adi]
    basliklar = [cell.value for cell in ws[1]]
    try:
        kimlik_col = basliklar.index("kimlik_no") + 1
        satir_col = basliklar.index("excel_satir") + 1
    except ValueError:
        return set()

    kullanilan = set()
    for row in range(2, ws.max_row + 1):
        if str(ws.cell(row=row, column=kimlik_col).value or "").strip() != str(kimlik_no).strip():
            continue
        try:
            kullanilan.add(int(ws.cell(row=row, column=satir_col).value))
        except (TypeError, ValueError):
            pass
    return kullanilan


def log_kaydi_ekle(wb, log_yolu: Path, sayfa_adi: str, kayit: Dict[str, object]) -> None:
    ws = wb[sayfa_adi]
    basliklar = LOG_SAYFA_BASLIKLARI[sayfa_adi]
    ws.append([kayit.get(baslik, "") for baslik in basliklar])
    wb.save(log_yolu)


# =========================
# WORD ALANLARINI DOLDURMA
# =========================

def hitap_alani_doldur(doc: Document, hitap: str) -> None:
    satirlar = hitap_satirlarini_ayir(hitap)

    # Şablonda İlgi paragrafından önceki sarı paragraflar hitap alanlarıdır.
    hitap_paragraflari = []
    for p in doc.paragraphs:
        if p.text.strip().startswith("İlgi:"):
            break
        if p.text.strip() and paragrafta_sari_var_mi(p):
            hitap_paragraflari.append(p)

    if not hitap_paragraflari:
        raise ValueError("Word şablonunda hitap için sarı alan bulunamadı.")

    # 1. satır: örn. ADANA DEFTERDARLIĞI
    fmt1 = paragraf_ilk_run_format(hitap_paragraflari[0])
    paragraf_yaz(
        hitap_paragraflari[0],
        turkce_buyuk_harf(satirlar[0]),
        fmt=fmt1,
        sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE,
    )

    # 2. satır ve varsa devamı: örn. (Ziyapaşa Vergi Dairesi Müdürlüğü)
    if len(hitap_paragraflari) >= 2:
        ikinci_metin = "\n".join(paranteze_al(satir) for satir in satirlar[1:]) if len(satirlar) > 1 else ""
        fmt2 = paragraf_ilk_run_format(hitap_paragraflari[1])
        paragraf_yaz(
            hitap_paragraflari[1],
            ikinci_metin,
            fmt=fmt2,
            sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE,
        )
    elif len(satirlar) > 1:
        # Şablonda tek hitap paragrafı varsa kalan satırları aynı paragrafa alt satır olarak ekler.
        paragraf_yaz(
            hitap_paragraflari[0],
            "\n".join([turkce_buyuk_harf(satirlar[0]), *[paranteze_al(satir) for satir in satirlar[1:]]]),
            fmt=fmt1,
            sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE,
        )


def konu_alani_doldur(doc: Document, ad_soyad: str) -> None:
    yeni_konu = f"{ad_soyad} Hk."

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(text == "Konu" for text in row_text):
                hedef_cell = None
                for cell in row.cells:
                    if "Hk" in cell.text or paragrafta_sari_var_mi(hucre_ilk_paragraf(cell)):
                        hedef_cell = cell
                        break
                if hedef_cell is None and len(row.cells) >= 3:
                    hedef_cell = row.cells[2]
                if hedef_cell is None:
                    raise ValueError("Konu satırı bulundu fakat konu hücresi belirlenemedi.")

                p = hucre_ilk_paragraf(hedef_cell)
                fmt = paragraf_ilk_run_format(p)
                paragraf_yaz(p, yeni_konu, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
                return

    raise ValueError("Word şablonunda 'Konu' satırı bulunamadı.")


def ilgi_alani_doldur(doc: Document, ilgi_tarihi: str, ilgi_sayisi: str) -> None:
    yeni_ilgi = f"İlgi: {ilgi_tarihi} tarihli ve {ilgi_sayisi} sayılı yazınız."
    for p in doc.paragraphs:
        if p.text.strip().startswith("İlgi:"):
            fmt = paragraf_ilk_run_format(p)
            paragraf_yaz(p, yeni_ilgi, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
            return
    raise ValueError("Word şablonunda 'İlgi:' paragrafı bulunamadı.")


def ana_metin_alani_doldur(doc: Document, ad_soyad: str, kimlik_no: str, kisi_turu: str) -> None:
    kisi_turu_norm = turkce_buyuk_harf(kisi_turu)
    kimlik_etiketi = "VKN" if "TÜZEL" in kisi_turu_norm else "TCKN"
    yatirimci_ifadesi = "unvanlı yatırımcının" if "TÜZEL" in kisi_turu_norm else "adlı yatırımcının"

    yeni_metin = (
        f"İlgide kayıtlı yazınızda özetle, {ad_soyad} ({kimlik_etiketi}:{kimlik_no}) {yatirimci_ifadesi} "
        "vergi borçlarından dolayı Merkezimiz nezdinde yer alan hak ve alacaklarına "
        "6183 sayılı Amme Alacaklarının Tahsil Usulü Hakkında Kanun'un 79'uncu maddesi "
        "gereğince haciz konulduğu bildirilmiştir. "
    )

    for p in doc.paragraphs:
        if p.text.strip().startswith("İlgide kayıtlı yazınızda özetle"):
            fmt = paragraf_ilk_run_format(p)
            paragraf_yaz(p, yeni_metin, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)
            return
    raise ValueError("Word şablonunda ana metin paragrafı bulunamadı.")


def tum_sari_vurgulari_temizle(doc: Document) -> None:
    """Kalan sarı vurguları temizler. Dinamik alanlar dışında sarı not bırakılmadıysa faydalıdır."""
    def temizle_paragraphs(paragraphs: Iterable):
        for p in paragraphs:
            for run in p.runs:
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    run.font.highlight_color = None

    temizle_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                temizle_paragraphs(cell.paragraphs)

    for section in doc.sections:
        temizle_paragraphs(section.header.paragraphs)
        temizle_paragraphs(section.footer.paragraphs)
        for table in section.header.tables + section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    temizle_paragraphs(cell.paragraphs)


def word_olustur(sablon_yolu: Path, cikti_yolu: Path, veri: Dict[str, str]) -> None:
    doc = Document(sablon_yolu)

    hitap_alani_doldur(doc, veri["hitap"])
    konu_alani_doldur(doc, veri["ad_soyad"])
    ilgi_alani_doldur(doc, veri["ilgi_tarihi"], veri["ilgi_sayisi"])
    ana_metin_alani_doldur(doc, veri["ad_soyad"], veri["kimlik_no"], veri["kisi_turu"])

    if FINALDE_SARI_ALANLARI_TEMIZLE:
        tum_sari_vurgulari_temizle(doc)

    cikti_yolu.parent.mkdir(parents=True, exist_ok=True)
    doc.save(cikti_yolu)


def mahkeme_tereke_word_olustur(
    sablon_yolu: Path,
    cikti_yolu: Path,
    veri: Dict[str, str],
    bakiyeli_mi: bool,
) -> None:
    doc = Document(sablon_yolu)

    muris = veri["muris_ad_soyad"]
    kimlik_no = veri["kimlik_no"]
    kimlik_turu = veri["kimlik_turu"]
    dosya_tur = dosya_tur_metni(veri)
    ilgi_tarihi = veri["ilgi_tarihi"]

    hitap = mahkeme_hitabi_olustur(veri["mahkeme"])
    ilgi = f"İlgi: {ilgi_tarihi} tarihli ve {dosya_tur} sayılı yazınız."
    konu = f"{dosya_tur} ({muris}) Hk."

    if bakiyeli_mi:
        alan_degerleri = [
            hitap,
            ilgi,
            (
                f"İlgide kayıtlı yazınızda Mahkemenizde görülmekte dava kapsamında {muris} "
                f"({kimlik_turu}:{kimlik_no}) adlı murisin Merkezimiz uhdesindeki hesabına ilişkin bilgi ve "
                "belgelerin gönderilmesi talep edilmiştir. "
            ),
            (
                f"İlgi’de kayıtlı yazınızda yer alan {kimlik_turu} bilgisi kullanılarak Merkezimiz kayıtları "
                "üzerinde yapılan inceleme neticesinde, SPK'nun 83’üncü maddesinin dördüncü fıkrası "
                f"kapsamında {muris} adlı murisin Merkezimiz uhdesinde yer alan hesabında bulunan sermaye "
                "piyasası araçları ile sermaye piyasası araçlarının yazımız tarihi itibariyle değerlerini "
                "gösteren bakiye raporu (Ek.1) ile raporda yer alan bilgilere ilişkin açıklamalara "
                "(Ek.2) ekte yer verilmiştir. "
            ),
            f"EK: 1-Bakiye Raporu ({muris})",
            konu,
        ]
    else:
        alan_degerleri = [
            hitap,
            ilgi,
            (
                f"İlgide kayıtlı yazınız ile {muris} ({kimlik_turu}:{kimlik_no}) adlı murisin terekesinin iflas "
                "hükümlerine göre tasfiyesine karar verildiği ve tasfiye işlemlerine Mahkemenizin "
                f"{dosya_tur} sayılı dosyasından başlandığı TMK 622, İİK 226,180,181,184 ve "
                "208.maddeleri gereğince tebliğ edilmiştir."
            ),
            f"EK: Bakiye Raporu ({muris})",
            konu,
        ]

    sari_paragraflar = sari_paragraflari_bul(doc)
    if len(sari_paragraflar) < len(alan_degerleri):
        raise ValueError(
            f"Word şablonunda yeterli sarı alan bulunamadı. Beklenen: {len(alan_degerleri)}, bulunan: {len(sari_paragraflar)}"
        )

    for p, deger in zip(sari_paragraflar, alan_degerleri):
        fmt = paragraf_ilk_run_format(p)
        paragraf_yaz(p, deger, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)

    if FINALDE_SARI_ALANLARI_TEMIZLE:
        tum_sari_vurgulari_temizle(doc)

    cikti_yolu.parent.mkdir(parents=True, exist_ok=True)
    doc.save(cikti_yolu)


def takasbank_bildirim_satiri(veri: Dict[str, str]) -> str:
    hitap_satirlari = hitap_satirlarini_ayir(veri["hitap"])
    haczi_koyan = hitap_satirlari[1] if len(hitap_satirlari) > 1 else hitap_satirlari[0]
    return f"{haczi_koyan}’nün {veri['ilgi_tarihi']} tarihli ve {veri['ilgi_sayisi']} sayılı yazısı."


def takasbank_bildirim_word_olustur(
    sablon_yolu: Path,
    cikti_yolu: Path,
    veriler: List[Dict[str, str]],
) -> None:
    doc = Document(sablon_yolu)
    sari_paragraflar = sari_paragraflari_bul(doc)

    if len(sari_paragraflar) < 5:
        raise ValueError(f"Takasbank bildirim şablonunda 5 sarı liste alanı bekleniyor, bulunan: {len(sari_paragraflar)}")

    for i, p in enumerate(sari_paragraflar[:5]):
        fmt = paragraf_ilk_run_format(p)
        metin = takasbank_bildirim_satiri(veriler[i]) if i < len(veriler) else ""
        paragraf_yaz(p, metin, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)

    if FINALDE_SARI_ALANLARI_TEMIZLE:
        tum_sari_vurgulari_temizle(doc)

    cikti_yolu.parent.mkdir(parents=True, exist_ok=True)
    doc.save(cikti_yolu)


def ytm_ici_talep_ek_satiri(veri: Dict[str, str], sira_no: int) -> str:
    return (
        f"{sira_no}- {veri['mahkeme']}’nin {veri['ilgi_tarihi']} tarih ve "
        f"{dosya_tur_metni(veri)} sayılı yazısı."
    )


def ytm_ici_talep_tablosu_bul(doc: Document):
    for table in doc.tables:
        if len(table.columns) < 6 or len(table.rows) < 2:
            continue
        basliklar = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "adı soyad" in basliklar[0] and "tckn" in basliklar[1]:
            return table
    raise ValueError("YTM içi talep şablonunda 6 kolonlu yatırımcı tablosu bulunamadı.")


def ytm_ici_talep_konu_doldur(doc: Document, veriler: List[Dict[str, str]]) -> None:
    konu = "Bilgi Talebi (" + ", ".join(veri["muris_ad_soyad"] for veri in veriler) + ")"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Bilgi Talebi" in cell.text and any(paragrafta_sari_var_mi(p) for p in cell.paragraphs):
                    hucre_yaz(cell, konu)
                    return
    raise ValueError("YTM içi talep şablonunda sarı konu alanı bulunamadı.")


def ytm_ici_talep_tablosu_doldur(doc: Document, veriler: List[Dict[str, str]]) -> None:
    table = ytm_ici_talep_tablosu_bul(doc)
    if len(table.rows) < 2:
        raise ValueError("YTM içi talep şablonunda örnek veri satırı bulunamadı.")

    ornek_satir = deepcopy(table.rows[1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)

    for sira_no, veri in enumerate(veriler, start=1):
        yeni_satir = deepcopy(ornek_satir)
        table._tbl.append(yeni_satir)
        row = table.rows[-1]
        degerler = [
            veri["muris_ad_soyad"],
            veri["kimlik_no"],
            veri["mahkeme"],
            veri["ilgi_tarihi"],
            dosya_tur_metni(veri),
            str(sira_no),
        ]
        for cell, deger in zip(row.cells, degerler):
            hucre_yaz(cell, deger)


def ytm_ici_talep_ekleri_doldur(doc: Document, veriler: List[Dict[str, str]]) -> None:
    ek_paragraflar = [
        p
        for p in doc.paragraphs
        if paragrafta_sari_var_mi(p) and re.match(r"^\s*\d+\s*-", p.text.strip())
    ]
    if not ek_paragraflar:
        raise ValueError("YTM içi talep şablonunda sarı ek listesi alanı bulunamadı.")

    ek_metni = "\n".join(ytm_ici_talep_ek_satiri(veri, i) for i, veri in enumerate(veriler, start=1))
    fmt = paragraf_ilk_run_format(ek_paragraflar[0])
    paragraf_yaz(ek_paragraflar[0], ek_metni, fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)

    for p in ek_paragraflar[1:]:
        fmt = paragraf_ilk_run_format(p)
        paragraf_yaz(p, "", fmt=fmt, sariyi_temizle=FINALDE_SARI_ALANLARI_TEMIZLE)


def ytm_ici_talep_word_olustur(
    sablon_yolu: Path,
    cikti_yolu: Path,
    veriler: List[Dict[str, str]],
) -> None:
    if not veriler:
        raise ValueError("YTM içi talep yazısı için en az bir kayıt gerekli.")

    doc = Document(sablon_yolu)
    ytm_ici_talep_konu_doldur(doc, veriler)
    ytm_ici_talep_tablosu_doldur(doc, veriler)
    ytm_ici_talep_ekleri_doldur(doc, veriler)

    if FINALDE_SARI_ALANLARI_TEMIZLE:
        tum_sari_vurgulari_temizle(doc)

    cikti_yolu.parent.mkdir(parents=True, exist_ok=True)
    doc.save(cikti_yolu)


# =========================
# ANA AKIŞ
# =========================

def main() -> None:
    base_dir = uygulama_klasoru()

    root = tk.Tk()
    root.withdraw()

    try:
        secim = belge_turu_sec(root)
        if secim is None:
            return

        ayar = BELGE_TURU_AYARLARI[secim]
        if not ayar.aktif:
            messagebox.showinfo("Henüz tanımlı değil", ayar.aciklama, parent=root)
            return

        belge_islem = secim
        if secim == "MAHKEME_TEREKE":
            alt_secim = mahkeme_tereke_alt_sec(root)
            if alt_secim is None:
                return
            belge_islem = alt_secim

        tcknler = tcknleri_al(root)
        if tcknler is None:
            return

        excel_yolu = dosya_yoksa_sec(
            root,
            base_dir / EXCEL_DOSYA_ADI,
            "Excel dosyasını seçiniz",
            [("Excel dosyası", "*.xlsx")],
        )
        if excel_yolu is None:
            return

        if belge_islem == "MAHKEME_TEREKE_BAKIYELI":
            sablon_adi = MAHKEME_TEREKE_BAKIYELI_SABLON_ADI
        elif belge_islem == "MAHKEME_TEREKE_BAKIYESIZ":
            sablon_adi = MAHKEME_TEREKE_BAKIYESIZ_SABLON_ADI
        elif belge_islem == "TAKASBANK_BILDIRIM":
            sablon_adi = TAKASBANK_BILDIRIM_SABLON_ADI
        elif belge_islem == "YTM_ICI_TALEP":
            sablon_adi = YTM_ICI_TALEP_SABLON_ADI
        else:
            sablon_adi = WORD_SABLON_ADI

        sablon_yolu = dosya_yoksa_sec(
            root,
            base_dir / SABLON_KLASOR_ADI / sablon_adi,
            "Word şablon dosyasını seçiniz",
            [("Word dosyası", "*.docx")],
        )
        if sablon_yolu is None:
            return

        log_yolu = log_dosya_yolu(base_dir)
        log_wb = log_workbook_hazirla(log_yolu)

        olusan_dosyalar: List[Path] = []
        hatalar: List[str] = []
        uyarilar: List[str] = []

        def coklu_satir_uyarisi_ekle(veri: Dict[str, str], kimlik_anahtari: str) -> None:
            onceki_satirlar = veri.get("onceki_excel_satirlari", "")
            if onceki_satirlar:
                uyarilar.append(
                    f"{veri[kimlik_anahtari]} {veri['excel_sayfa']} sayfasında birden fazla satırda bulundu. "
                    f"{veri['excel_satir']}. satır kullanıldı; önceki satırlar: {onceki_satirlar}."
                )

        if belge_islem == "TAKASBANK_BILDIRIM":
            takasbank_verileri: List[Dict[str, str]] = []
            for tckn in tcknler:
                try:
                    veri = excelden_6183_kaydi_bul(excel_yolu, tckn)
                    coklu_satir_uyarisi_ekle(veri, "kimlik_no")
                    takasbank_verileri.append(veri)
                except Exception as exc:
                    hatalar.append(f"{tckn}: {exc}")

            for grup_no, baslangic in enumerate(range(0, len(takasbank_verileri), 5), start=1):
                grup = takasbank_verileri[baslangic:baslangic + 5]
                if not grup:
                    continue
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                dosya_adi = f"takasbank_bildirim_{timestamp}"
                if len(takasbank_verileri) > 5:
                    dosya_adi += f"_bolum_{grup_no}"
                cikti_yolu = base_dir / CIKTI_KLASOR_ADI / f"{dosya_adi}.docx"
                try:
                    takasbank_bildirim_word_olustur(sablon_yolu, cikti_yolu, grup)
                    olusan_dosyalar.append(cikti_yolu)
                    for veri in grup:
                        log_kaydi_ekle(
                            log_wb,
                            log_yolu,
                            "takasbank",
                            {
                                "log_zamani": datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                                "yazi_turu": "Takasbank Bildirim yazısı",
                                "kimlik_turu": kimlik_turu_belirle(veri["kimlik_no"]),
                                "kimlik_no": veri["kimlik_no"],
                                "kisi_turu": veri["kisi_turu"],
                                "ad_unvan": veri["ad_soyad"],
                                "excel_sayfa": veri["excel_sayfa"],
                                "excel_satir": veri["excel_satir"],
                                "sablon": sablon_adi,
                                "olusan_dosya": str(cikti_yolu),
                                "grup_no": grup_no,
                                "haczi_koyan": hitap_satirlarini_ayir(veri["hitap"])[1] if len(hitap_satirlarini_ayir(veri["hitap"])) > 1 else hitap_satirlarini_ayir(veri["hitap"])[0],
                                "haciz_bildirim_tarihi": veri["ilgi_tarihi"],
                                "haciz_bildirim_sayisi": veri["ilgi_sayisi"],
                            },
                        )
                except Exception as exc:
                    hatalar.append(f"Takasbank bildirim bölüm {grup_no}: {exc}")

        elif belge_islem == "YTM_ICI_TALEP":
            ytm_verileri: List[Dict[str, str]] = []
            for tckn in tcknler:
                try:
                    veri = excelden_tereke_kaydi_bul(excel_yolu, tckn)
                    coklu_satir_uyarisi_ekle(veri, "kimlik_no")
                    ytm_verileri.append(veri)
                except Exception as exc:
                    hatalar.append(f"{tckn}: {exc}")

            if ytm_verileri:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                dosya_adi = f"ytm_ici_talep_{timestamp}.docx"
                cikti_yolu = base_dir / CIKTI_KLASOR_ADI / dosya_adi
                try:
                    ytm_ici_talep_word_olustur(sablon_yolu, cikti_yolu, ytm_verileri)
                    olusan_dosyalar.append(cikti_yolu)
                    for sira_no, veri in enumerate(ytm_verileri, start=1):
                        log_kaydi_ekle(
                            log_wb,
                            log_yolu,
                            "ytm_ici_talep",
                            {
                                "log_zamani": datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                                "yazi_turu": "YTM İçi Talep Yazısı",
                                "kimlik_turu": veri["kimlik_turu"],
                                "kimlik_no": veri["kimlik_no"],
                                "muris_ad_soyad": veri["muris_ad_soyad"],
                                "excel_sayfa": veri["excel_sayfa"],
                                "excel_satir": veri["excel_satir"],
                                "sablon": sablon_adi,
                                "olusan_dosya": str(cikti_yolu),
                                "sira_no": sira_no,
                                "mahkeme": veri["mahkeme"],
                                "ilgi_tarihi": veri["ilgi_tarihi"],
                                "dosya_no": veri["dosya_no"],
                                "tur": veri["tur"],
                            },
                        )
                except Exception as exc:
                    hatalar.append(f"YTM İçi Talep Yazısı: {exc}")

        else:
            for tckn in tcknler:
                try:
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

                    if belge_islem == "MAHKEME_TEREKE_BAKIYELI":
                        veri = excelden_tereke_kaydi_bul(excel_yolu, tckn)
                        coklu_satir_uyarisi_ekle(veri, "tckn")
                        dosya_adi = f"{temiz_dosya_adi(veri['muris_ad_soyad'])}_{veri['tckn']}_tereke_bakiyeli_{timestamp}.docx"
                        cikti_yolu = base_dir / CIKTI_KLASOR_ADI / dosya_adi
                        mahkeme_tereke_word_olustur(sablon_yolu, cikti_yolu, veri, bakiyeli_mi=True)
                        log_kaydi_ekle(
                            log_wb,
                            log_yolu,
                            "mahkeme_tereke",
                            {
                                "log_zamani": datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                                "yazi_turu": "Mahkeme tereke yazısı",
                                "alt_tur": "Bakiyesi Var",
                                "kimlik_turu": veri["kimlik_turu"],
                                "kimlik_no": veri["tckn"],
                                "muris_ad_soyad": veri["muris_ad_soyad"],
                                "excel_sayfa": veri["excel_sayfa"],
                                "excel_satir": veri["excel_satir"],
                                "sablon": sablon_adi,
                                "olusan_dosya": str(cikti_yolu),
                                "mahkeme": veri["mahkeme"],
                                "dosya_no": veri["dosya_no"],
                                "tur": veri["tur"],
                            },
                        )
                    elif belge_islem == "MAHKEME_TEREKE_BAKIYESIZ":
                        veri = excelden_tereke_kaydi_bul(excel_yolu, tckn)
                        coklu_satir_uyarisi_ekle(veri, "tckn")
                        dosya_adi = f"{temiz_dosya_adi(veri['muris_ad_soyad'])}_{veri['tckn']}_tereke_bakiyesiz_{timestamp}.docx"
                        cikti_yolu = base_dir / CIKTI_KLASOR_ADI / dosya_adi
                        mahkeme_tereke_word_olustur(sablon_yolu, cikti_yolu, veri, bakiyeli_mi=False)
                        log_kaydi_ekle(
                            log_wb,
                            log_yolu,
                            "mahkeme_tereke",
                            {
                                "log_zamani": datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                                "yazi_turu": "Mahkeme tereke yazısı",
                                "alt_tur": "Bakiyesi yok",
                                "kimlik_turu": veri["kimlik_turu"],
                                "kimlik_no": veri["tckn"],
                                "muris_ad_soyad": veri["muris_ad_soyad"],
                                "excel_sayfa": veri["excel_sayfa"],
                                "excel_satir": veri["excel_satir"],
                                "sablon": sablon_adi,
                                "olusan_dosya": str(cikti_yolu),
                                "mahkeme": veri["mahkeme"],
                                "dosya_no": veri["dosya_no"],
                                "tur": veri["tur"],
                            },
                        )
                    else:
                        veri = excelden_6183_kaydi_bul(excel_yolu, tckn)
                        coklu_satir_uyarisi_ekle(veri, "kimlik_no")
                        dosya_adi = f"{temiz_dosya_adi(veri['ad_soyad'])}_{veri['kimlik_no']}_6183_{timestamp}.docx"
                        cikti_yolu = base_dir / CIKTI_KLASOR_ADI / dosya_adi
                        word_olustur(sablon_yolu, cikti_yolu, veri)
                        log_kaydi_ekle(
                            log_wb,
                            log_yolu,
                            "6183",
                            {
                                "log_zamani": datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                                "yazi_turu": "6183 cevap yazısı",
                                "kimlik_turu": kimlik_turu_belirle(veri["kimlik_no"]),
                                "kimlik_no": veri["kimlik_no"],
                                "kisi_turu": veri["kisi_turu"],
                                "ad_unvan": veri["ad_soyad"],
                                "excel_sayfa": veri["excel_sayfa"],
                                "excel_satir": veri["excel_satir"],
                                "sablon": sablon_adi,
                                "olusan_dosya": str(cikti_yolu),
                                "hitap": veri["hitap"],
                                "ilgi_tarihi": veri["ilgi_tarihi"],
                                "ilgi_sayisi": veri["ilgi_sayisi"],
                            },
                        )

                    olusan_dosyalar.append(cikti_yolu)
                except Exception as exc:
                    hatalar.append(f"{tckn}: {exc}")

        mesaj = [f"Oluşturulan Word dosyası sayısı: {len(olusan_dosyalar)}"]
        if olusan_dosyalar:
            mesaj.append("")
            mesaj.extend(str(yol) for yol in olusan_dosyalar)
        if uyarilar:
            mesaj.append("")
            mesaj.append("Uyarılar:")
            mesaj.extend(uyarilar)
        if hatalar:
            mesaj.append("")
            mesaj.append("Oluşturulamayan kayıtlar:")
            mesaj.extend(hatalar)

        messagebox.showinfo("İşlem tamamlandı", "\n".join(mesaj), parent=root)

        # Windows'ta tek dosya oluştuysa dosyayı, birden fazla dosya oluştuysa çıktı klasörünü açar.
        try:
            if os.name == "nt" and olusan_dosyalar:
                acilacak_yol = olusan_dosyalar[0] if len(olusan_dosyalar) == 1 else olusan_dosyalar[0].parent
                os.startfile(str(acilacak_yol))  # type: ignore[attr-defined]
        except Exception:
            pass

    except Exception as exc:
        messagebox.showerror("Hata", str(exc), parent=root)
        raise
    finally:
        try:
            root.destroy()
        except Exception:
            pass


if __name__ == "__main__":
    main()
