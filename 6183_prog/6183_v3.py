import os
import re
from pathlib import Path
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from typing import Dict, List, Any, Tuple, Optional

import fitz  # PyMuPDF

import tkinter as tk
from tkinter import filedialog, messagebox

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ============================================================
# Regex ve yardımcılar
# ============================================================

DATE_RE = re.compile(r"\b\d{2}/\d{2}/\d{4}\b")
AMOUNT_RE = re.compile(r"^\d{1,3}(?:\.\d{3})*,\d+$|^\d+,\d+$")

TR_LOWER_MAP = str.maketrans({
    "I": "ı",
    "İ": "i",
    "Ğ": "ğ",
    "Ü": "ü",
    "Ş": "ş",
    "Ö": "ö",
    "Ç": "ç",
})

TR_UPPER_MAP = {
    "ı": "I",
    "i": "İ",
    "ğ": "Ğ",
    "ü": "Ü",
    "ş": "Ş",
    "ö": "Ö",
    "ç": "Ç",
}

TURKISH_ASCII_MAP = str.maketrans({
    "İ": "I",
    "I": "I",
    "ı": "i",
    "Ğ": "G",
    "ğ": "g",
    "Ü": "U",
    "ü": "u",
    "Ş": "S",
    "ş": "s",
    "Ö": "O",
    "ö": "o",
    "Ç": "C",
    "ç": "c",
})


def tr_lower(text: str) -> str:
    return text.translate(TR_LOWER_MAP).lower()


def tr_title_word(word: str) -> str:
    word = tr_lower(word.strip())
    if not word:
        return word
    first = TR_UPPER_MAP.get(word[0], word[0].upper())
    return first + word[1:]


def format_investor_name(name: str) -> str:
    """
    TÜRKER ATEŞ   -> Türker ATEŞ
    MUSTAFA ALTAŞ -> Mustafa ALTAŞ
    """
    parts = [p for p in name.split() if p.strip()]

    if len(parts) >= 2:
        first_names = " ".join(tr_title_word(p) for p in parts[:-1])
        surname = parts[-1].upper()
        return f"{first_names} {surname}"

    return tr_title_word(name)


def normalize_text(text: str) -> str:
    text = str(text).replace("\xa0", " ").strip()
    text = text.translate(TURKISH_ASCII_MAP).lower()
    text = re.sub(r"[^a-z0-9/]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def parse_decimal_tr(value: str) -> Decimal:
    """
    109.405,71  -> Decimal('109405.71')
    690.000,000 -> Decimal('690000.000')
    """
    value = str(value).strip()
    if not value:
        return Decimal("0")
    value = value.replace(".", "").replace(",", ".")
    return Decimal(value)


def format_decimal_tr(value: Decimal, decimals: int) -> str:
    """
    Decimal('109405.71') -> 109.405,71
    Decimal('690000.000') -> 690.000,000
    """
    q = Decimal("1").scaleb(-decimals)
    value = value.quantize(q, rounding=ROUND_HALF_UP)

    formatted = f"{value:,.{decimals}f}"
    formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
    return formatted


def tr_number_to_float(value: str) -> float:
    try:
        return float(str(value).replace(".", "").replace(",", "."))
    except Exception:
        return 0.0


# ============================================================
# PDF satır ve alan okuma
# ============================================================

HEADER_PATTERNS = [
    ("sahsi_musterek", ["sahsi/musterek"]),
    ("hesap_orani", ["hesap", "oran"]),
    ("acilis_tarih", ["acilis", "tarih"]),
    ("borsada_islem", ["borsada", "islem"]),

    ("mkk_sicil_no", ["mkk", "sicil", "no"]),
    ("tvs_sicil_no", ["tvs", "sicil", "no"]),
    ("kimlik_no", ["kimlik", "no"]),
    ("musteri_no", ["musteri", "no"]),

    ("hesap_tipi", ["hesap", "tipi"]),
    ("tanim_adi", ["tanim", "adi"]),
    ("ek_tanim", ["ek", "tanim"]),
    ("bakiye_tipi", ["bakiye", "tipi"]),

    ("ad", ["ad"]),
    ("soyad", ["soyad"]),
    ("bakiye", ["bakiye"]),
    ("tutar", ["tutar"]),
    ("tanim", ["tanim"]),
    ("grup", ["grup"]),
]


def is_rotated_page(page, words) -> bool:
    if page.rotation in (90, 270):
        return True

    widths = [abs(w[2] - w[0]) for w in words if str(w[4]).strip()]
    heights = [abs(w[3] - w[1]) for w in words if str(w[4]).strip()]

    if not widths or not heights:
        return False

    widths.sort()
    heights.sort()
    median_width = widths[len(widths) // 2]
    median_height = heights[len(heights) // 2]

    return median_height > median_width * 1.35


def group_words_by_row(words, rotated: bool) -> List[List[Tuple]]:
    """
    Döndürülmüş raporlarda gerçek satırlar X koordinatına göre gruplanır.
    Normal PDF'lerde gerçek satırlar Y koordinatına göre gruplanır.
    """
    if rotated:
        key_func = lambda w: float(w[0])
    else:
        key_func = lambda w: float(w[1])

    sorted_words = sorted(words, key=key_func)
    groups: List[Dict[str, Any]] = []

    for w in sorted_words:
        key = key_func(w)

        if not groups or abs(key - groups[-1]["key"]) > 3.0:
            groups.append({"key": key, "words": [w]})
        else:
            groups[-1]["words"].append(w)
            groups[-1]["key"] = sum(key_func(x) for x in groups[-1]["words"]) / len(groups[-1]["words"])

    return [g["words"] for g in groups]


def axis_span(word, rotated: bool) -> Tuple[float, float]:
    """
    Döndürülmüş sayfada kolon ekseni Y'dir.
    Normal sayfada kolon ekseni X'tir.
    """
    if rotated:
        return float(word[1]), float(word[3])
    return float(word[0]), float(word[2])


def axis_center(obj: Dict[str, Any]) -> float:
    return (obj["start"] + obj["end"]) / 2


def overlap(a_start: float, a_end: float, b_start: float, b_end: float) -> float:
    return max(0.0, min(a_end, b_end) - max(a_start, b_start))


def words_to_data_cells(row_words, rotated: bool) -> List[Dict[str, Any]]:
    """
    Veri satırındaki hücreleri çıkarır.
    PyMuPDF word tuple: x0, y0, x1, y1, text, block_no, line_no, word_no
    """
    cells = []

    line_numbers = sorted(set(int(w[6]) for w in row_words))

    for ln in line_numbers:
        parts = [w for w in row_words if int(w[6]) == ln]
        parts = sorted(parts, key=lambda w: int(w[7]))

        text = " ".join(str(w[4]).strip() for w in parts if str(w[4]).strip())

        if not text:
            continue

        starts = []
        ends = []

        for w in parts:
            s, e = axis_span(w, rotated)
            starts.append(s)
            ends.append(e)

        cells.append({
            "text": text,
            "start": min(starts),
            "end": max(ends),
            "line_no": ln,
        })

    return cells


def split_cells_by_line(data_cells: List[Dict[str, Any]]) -> List[List[Dict[str, Any]]]:
    grouped: Dict[int, List[Dict[str, Any]]] = {}

    for cell in data_cells:
        grouped.setdefault(int(cell["line_no"]), []).append(cell)

    return [grouped[line_no] for line_no in sorted(grouped)]


def find_pattern_spans(header_words, rotated: bool) -> Dict[str, Dict[str, Any]]:
    """
    Başlık satırındaki alan adlarını koordinatlarıyla yakalar.

    Önemli:
    'Hesap Oran Açılış Tarih' tek line_no içinde gelse bile
    'Hesap Oran' ve 'Açılış Tarih' ayrı ayrı yakalanır.
    """
    if rotated:
        ordered = sorted(header_words, key=lambda w: -float(w[1]))
    else:
        ordered = sorted(header_words, key=lambda w: float(w[0]))

    tokens = [normalize_text(w[4]) for w in ordered]
    used = set()
    spans: Dict[str, Dict[str, Any]] = {}

    for field_name, pattern in HEADER_PATTERNS:
        pattern_len = len(pattern)

        for i in range(0, len(tokens) - pattern_len + 1):
            indexes = list(range(i, i + pattern_len))

            if any(idx in used for idx in indexes):
                continue

            if tokens[i:i + pattern_len] == pattern:
                starts = []
                ends = []

                for idx in indexes:
                    s, e = axis_span(ordered[idx], rotated)
                    starts.append(s)
                    ends.append(e)
                    used.add(idx)

                spans[field_name] = {
                    "field": field_name,
                    "text": " ".join(str(ordered[idx][4]) for idx in indexes),
                    "start": min(starts),
                    "end": max(ends),
                }
                break

    return spans


def row_has_table_header(row_words) -> bool:
    text = " ".join(str(w[4]) for w in row_words)
    norm = normalize_text(text)

    return (
        "ad" in norm
        and "musteri" in norm
        and "bakiye" in norm
        and "tutar" in norm
        and "tanim" in norm
    )


def map_data_by_header_spans(
    header_spans: Dict[str, Dict[str, Any]],
    data_cells: List[Dict[str, Any]]
) -> Dict[str, str]:
    """
    Alan adı koordinat aralığına en iyi uyan veri hücresini bulur.
    Bu sayede Şahsi/Müşterek = S ve Hesap Oran = 100 doğru eşleşir.
    """
    result: Dict[str, str] = {}
    used_cells = set()

    # Önce uzun/kritik alanları eşleştirelim.
    field_order = [
        "ad", "soyad", "kimlik_no", "sahsi_musterek", "hesap_orani",
        "acilis_tarih", "borsada_islem", "mkk_sicil_no", "tvs_sicil_no",
        "musteri_no", "bakiye", "tutar", "hesap_tipi", "tanim",
        "tanim_adi", "grup", "ek_tanim", "bakiye_tipi",
    ]

    for field in field_order:
        if field not in header_spans:
            continue

        h = header_spans[field]
        h_center = axis_center(h)

        best_index = None
        best_score = None

        for i, cell in enumerate(data_cells):
            if i in used_cells:
                continue

            ov = overlap(h["start"], h["end"], cell["start"], cell["end"])
            distance = abs(h_center - axis_center(cell))

            # Öncelik çakışma, sonra yakınlık.
            score = ov * 1000 - distance

            if best_score is None or score > best_score:
                best_score = score
                best_index = i

        if best_index is not None:
            result[field] = data_cells[best_index]["text"]
            used_cells.add(best_index)

    return result


def parse_data_cells_fallback(data_cells: List[Dict[str, Any]]) -> Dict[str, str]:
    texts = [str(cell.get("text", "")).strip() for cell in data_cells if str(cell.get("text", "")).strip()]

    if len(texts) < 8:
        return {}

    mapped: Dict[str, str] = {}

    remaining = texts[:]

    if remaining and remaining[-1] in {"E", "H"}:
        mapped["borsada_islem"] = remaining.pop()

    if remaining and DATE_RE.search(remaining[-1]):
        mapped["acilis_tarih"] = remaining.pop()

    if remaining and remaining[-1].upper() in {"S", "M"}:
        mapped["sahsi_musterek"] = remaining.pop()

    if remaining and re.fullmatch(r"%?\d+(?:[.,]\d+)?", remaining[-1]):
        mapped["hesap_orani"] = remaining.pop()

    if remaining:
        mapped["kimlik_no"] = remaining.pop()
    if remaining:
        mapped["tvs_sicil_no"] = remaining.pop()
    if remaining:
        mapped["mkk_sicil_no"] = remaining.pop()
    if remaining:
        mapped["soyad"] = remaining.pop()

    if len(remaining) < 3:
        return mapped

    mapped["ad"] = remaining.pop(0)
    mapped["musteri_no"] = remaining.pop(0)
    mapped["bakiye"] = remaining.pop(0)

    if remaining:
        mapped["tanim"] = remaining.pop(0)

    if remaining:
        mapped["tanim_adi"] = remaining.pop(0)

    if remaining:
        mapped["hesap_tipi"] = remaining.pop(0)

    if remaining:
        mapped["grup"] = remaining.pop(0)

    if len(remaining) == 3:
        mapped["ek_tanim"] = remaining.pop(0)
        mapped["bakiye_tipi"] = remaining.pop(0)
        mapped["tutar"] = remaining.pop(0)
    elif len(remaining) == 2:
        if AMOUNT_RE.match(remaining[1]):
            mapped["bakiye_tipi"] = remaining[0]
            mapped["tutar"] = remaining[1]
        else:
            mapped["ek_tanim"] = remaining[0]
            mapped["bakiye_tipi"] = remaining[1]
    elif len(remaining) == 1:
        if AMOUNT_RE.match(remaining[0]):
            mapped["tutar"] = remaining[0]
        else:
            mapped["bakiye_tipi"] = remaining[0]

    return mapped


def extract_report_date(full_text: str) -> str:
    """
    Rapor tarihini alır.
    Genelde ikinci Tarih satırı bakiye tarihidir.
    """
    lines = [line.strip() for line in full_text.splitlines() if line.strip()]
    found = []

    for i, line in enumerate(lines):
        if normalize_text(line).startswith("tarih"):
            found.extend(DATE_RE.findall(line))
            if i + 1 < len(lines):
                found.extend(DATE_RE.findall(lines[i + 1]))

    if found:
        return found[-1]

    all_dates = DATE_RE.findall(full_text)
    return all_dates[-1] if all_dates else ""


def build_row_dict(mapped: Dict[str, str], report_date: str, page_no: int, raw_line: str) -> Dict[str, str]:
    ad = mapped.get("ad", "").strip()
    soyad = mapped.get("soyad", "").strip()
    adsoyad_raw = f"{ad} {soyad}".strip()

    return {
        "tarih": report_date,
        "ad": ad,
        "soyad": soyad,
        "adsoyad_raw": adsoyad_raw,
        "adsoyad": format_investor_name(adsoyad_raw),

        "kimlik_no": mapped.get("kimlik_no", "").strip(),
        "sahsi_musterek": mapped.get("sahsi_musterek", "").strip().upper(),
        "hesap_orani": mapped.get("hesap_orani", "").strip(),

        "acilis_tarih": mapped.get("acilis_tarih", "").strip(),
        "borsada_islem": mapped.get("borsada_islem", "").strip(),

        "mkk_sicil_no": mapped.get("mkk_sicil_no", "").strip(),
        "tvs_sicil_no": mapped.get("tvs_sicil_no", "").strip(),

        "musteri_no": mapped.get("musteri_no", "").strip(),
        "bakiye": mapped.get("bakiye", "").strip(),
        "tutar": mapped.get("tutar", "").strip(),
        "hesap_tipi": mapped.get("hesap_tipi", "").strip(),
        "tanim": mapped.get("tanim", "").strip(),
        "tanim_adi": mapped.get("tanim_adi", "").strip(),
        "grup": mapped.get("grup", "").strip(),
        "ek_tanim": mapped.get("ek_tanim", "").strip(),
        "bakiye_tipi": mapped.get("bakiye_tipi", "").strip(),

        "page": str(page_no),
        "raw_line": raw_line,
    }


def row_is_balance_row(row: Dict[str, str]) -> bool:
    if not row.get("adsoyad_raw"):
        return False

    if not row.get("musteri_no"):
        return False

    if not AMOUNT_RE.match(row.get("bakiye", "")):
        return False

    if row.get("tutar") and not AMOUNT_RE.match(row.get("tutar", "")):
        return False

    if not row.get("tanim"):
        return False

    return True


def extract_balance_rows(pdf_path: Path) -> Tuple[List[Dict[str, str]], List[str]]:
    rows: List[Dict[str, str]] = []
    debug_lines: List[str] = []

    with fitz.open(pdf_path) as doc:
        full_text_parts = []

        for page_no, page in enumerate(doc, start=1):
            full_text_parts.append(page.get_text("text") or "")

        full_text = "\n".join(full_text_parts)
        report_date = extract_report_date(full_text)

        current_header_spans: Optional[Dict[str, Dict[str, Any]]] = None

        for page_no, page in enumerate(doc, start=1):
            words = page.get_text("words")

            if not words:
                continue

            rotated = is_rotated_page(page, words)
            row_groups = group_words_by_row(words, rotated)

            for group in row_groups:
                raw_line = " | ".join(c["text"] for c in words_to_data_cells(group, rotated))
                debug_lines.append(raw_line)

                norm_line = normalize_text(raw_line)

                if "rapor sonu" in norm_line:
                    current_header_spans = None
                    continue

                if row_has_table_header(group):
                    current_header_spans = find_pattern_spans(group, rotated)
                    debug_lines.append(">>> BAŞLIK ALANLARI: " + ", ".join(sorted(current_header_spans.keys())))
                    continue

                if not current_header_spans:
                    continue

                data_cells = words_to_data_cells(group, rotated)
                candidate_cell_groups = [data_cells]
                split_groups = split_cells_by_line(data_cells)

                if len(split_groups) > 1 and any(len(cell_line) > 1 for cell_line in split_groups):
                    candidate_cell_groups = split_groups

                for cell_group in candidate_cell_groups:
                    candidate_raw_line = " | ".join(cell["text"] for cell in cell_group)
                    fallback_mapped = parse_data_cells_fallback(cell_group)
                    mapped = map_data_by_header_spans(current_header_spans, cell_group)

                    for key, value in fallback_mapped.items():
                        if not mapped.get(key):
                            mapped[key] = value

                    row = build_row_dict(mapped, report_date, page_no, candidate_raw_line or raw_line)

                    if row_is_balance_row(row):
                        rows.append(row)
                        continue

                    if fallback_mapped:
                        fallback_row = build_row_dict(fallback_mapped, report_date, page_no, candidate_raw_line or raw_line)
                        if row_is_balance_row(fallback_row):
                            rows.append(fallback_row)

    return rows, debug_lines


def has_positive_balance(row: Dict[str, str]) -> bool:
    bakiye = tr_number_to_float(row.get("bakiye", "0"))
    tutar = tr_number_to_float(row.get("tutar", "0"))
    return bakiye > 0 or tutar > 0


# ============================================================
# Word metni üretme
# ============================================================

def get_ratio_decimal(row: Dict[str, str]) -> Decimal:
    oran = str(row.get("hesap_orani", "")).strip()

    if not oran:
        return Decimal("100")

    oran = oran.replace("%", "").replace(",", ".")
    try:
        return Decimal(oran)
    except Exception:
        return Decimal("100")


def get_ratio_text(row: Dict[str, str]) -> str:
    oran = str(row.get("hesap_orani", "")).strip()

    if not oran:
        return "(Hesap Oranı: %100)"

    oran = oran.replace("%", "")
    return f"(Hesap Oranı: %{oran})"


def adjusted_values(row: Dict[str, str]) -> Tuple[str, str]:
    """
    S için değerler aynen yazılır.
    M için tutar ve bakiye Hesap Oranı / 100 ile çarpılır.
    """
    sahsi_musterek = row.get("sahsi_musterek", "").upper()
    ratio = get_ratio_decimal(row)

    tutar = parse_decimal_tr(row.get("tutar", "0"))
    bakiye = parse_decimal_tr(row.get("bakiye", "0"))

    if sahsi_musterek == "M":
        tutar = tutar * ratio / Decimal("100")
        bakiye = bakiye * ratio / Decimal("100")

    tutar_text = format_decimal_tr(tutar, 2)
    bakiye_text = format_decimal_tr(bakiye, 3)

    return tutar_text, bakiye_text


def format_tutar_for_text(row: Dict[str, str], tutar_text: str) -> str:
    raw_tutar = str(row.get("tutar", "")).strip()

    if not raw_tutar:
        return "0"

    try:
        if parse_decimal_tr(raw_tutar) == Decimal("0"):
            return "0"
    except Exception:
        pass

    return tutar_text


def is_zero_tutar(row: Dict[str, str]) -> bool:
    raw_tutar = str(row.get("tutar", "")).strip()

    if not raw_tutar:
        return True

    try:
        return parse_decimal_tr(raw_tutar) == Decimal("0")
    except Exception:
        return False


def account_type_text(row: Dict[str, str]) -> str:
    sahsi_musterek = row.get("sahsi_musterek", "").upper()

    if sahsi_musterek == "S":
        return "tekil hesabında"

    if sahsi_musterek == "M":
        return "müşterek hesabında"

    return "hesabında"


def group_label_text(grup: str) -> str:
    grup = str(grup).strip().upper()

    if grup == "E":
        return "hisse senedi"

    if grup == "F":
        return "yatırım fonu"

    return str(grup).strip()


def investment_phrase(row: Dict[str, str]) -> str:
    tutar_text, bakiye_text = adjusted_values(row)
    tutar_text = format_tutar_for_text(row, tutar_text)

    tanim = row.get("tanim", "")
    grup_code = str(row.get("grup", "")).strip().upper()
    grup = group_label_text(grup_code)

    if tanim.strip().upper() == "TL":
        return f"{tutar_text} TL değerinde nakit"

    if tanim and grup:
        if grup_code == "E" and is_zero_tutar(row):
            return f"{tutar_text} TL değerinde {bakiye_text} adet {tanim} tanımlı borsada işlem görmeyen hisse senedi"
        return f"{tutar_text} TL değerinde {bakiye_text} adet {tanim} tanımlı {grup}"

    if tanim:
        return f"{tutar_text} TL değerinde {bakiye_text} adet {tanim} tanımlı kıymeti"

    return f"{tutar_text} TL değerinde {bakiye_text} adet kıymeti"


def build_group_phrases(rows: List[Dict[str, str]]) -> List[str]:
    phrases: List[str] = []
    total_tl = Decimal("0")

    for row in rows:
        tanim = str(row.get("tanim", "")).strip().upper()
        if tanim == "TL":
            tutar_text, _ = adjusted_values(row)
            total_tl += parse_decimal_tr(tutar_text)
            continue

        phrases.append(investment_phrase(row))

    if total_tl > 0:
        phrases.insert(0, f"{format_decimal_tr(total_tl, 2)} TL değerinde nakit")

    return phrases


def join_phrases_turkish(phrases: List[str]) -> str:
    if not phrases:
        return ""

    if len(phrases) == 1:
        return phrases[0]

    if len(phrases) == 2:
        return f"{phrases[0]} ve {phrases[1]}"

    return "; ".join(phrases[:-1]) + f" ve {phrases[-1]}"


def format_investor_subject(name: str, kimlik_no: str) -> str:
    name = name.strip() if name else "İlgili yatırımcı"
    kimlik_no = kimlik_no.strip()

    if kimlik_no:
        return f"{name} (TCKN: {kimlik_no})"

    return name


def group_rows_by_investor(rows: List[Dict[str, str]]) -> Dict[Tuple[str, str, str, str, str, str], List[Dict[str, str]]]:
    grouped: Dict[Tuple[str, str, str, str, str, str], List[Dict[str, str]]] = {}

    for row in rows:
        key = (
            row.get("adsoyad", ""),
            row.get("kimlik_no", ""),
            row.get("musteri_no", ""),
            row.get("tarih", ""),
            row.get("sahsi_musterek", ""),
            row.get("hesap_orani", ""),
        )
        grouped.setdefault(key, []).append(row)

    return grouped


def build_word_paragraphs(rows: List[Dict[str, str]]) -> List[str]:
    grouped = group_rows_by_investor(rows)
    paragraphs = []

    for key, group_rows in grouped.items():
        first = group_rows[0]

        adsoyad = first.get("adsoyad", "")
        kimlik_no = first.get("kimlik_no", "")
        musteri_no = first.get("musteri_no", "")
        tarih = first.get("tarih", "")
        hesap_turu = account_type_text(first)
        oran_text = get_ratio_text(first)
        investor_subject = format_investor_subject(adsoyad, kimlik_no)

        phrases = build_group_phrases(group_rows)
        balance_text = join_phrases_turkish(phrases)

        paragraph = (
            f"SPK'nun 83’üncü maddesinin dördüncü fıkrası kapsamında "
            f"{investor_subject} adlı yatırımcının Merkezimiz uhdesindeki "
            f"{musteri_no} numaralı {hesap_turu} {oran_text} "
            f"{tarih} tarihi itibariyle {balance_text} bulunmaktadır."
        )

        paragraphs.append(paragraph)

    return paragraphs


def infer_name_from_filename(pdf_path: Path) -> str:
    stem = pdf_path.stem.replace("_", " ").replace("-", " ")
    cleaned = re.sub(r"\bpdfsam\b", " ", stem, flags=re.IGNORECASE)
    cleaned = re.sub(r"\bm[üu]zekkere\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\beki\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\b(?:hr|br)\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\d+", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    if not cleaned:
        return ""

    return " ".join(tr_title_word(part) for part in cleaned.split())


def extract_report_identity(pdf_path: Path) -> Dict[str, str]:
    with fitz.open(pdf_path) as doc:
        full_text = "\n".join(page.get_text("text") or "" for page in doc)

    report_date = extract_report_date(full_text)
    lines = [line.strip() for line in full_text.splitlines() if line.strip()]

    kimlik_no = ""
    for i, line in enumerate(lines):
        norm = normalize_text(line)
        if "tckn" in norm or "kimlik no" in norm:
            candidates = DATE_RE.sub("", line)
            digits = re.findall(r"\b\d{11}\b", candidates)
            if digits:
                kimlik_no = digits[-1]
                break
            if i + 1 < len(lines):
                digits = re.findall(r"\b\d{11}\b", lines[i + 1])
                if digits:
                    kimlik_no = digits[-1]
                    break

    return {
        "adsoyad": infer_name_from_filename(pdf_path),
        "kimlik_no": kimlik_no,
        "tarih": report_date,
    }


def build_empty_report_paragraphs(rows: List[Dict[str, str]], pdf_path: Path) -> List[str]:
    paragraphs: List[str] = []

    if rows:
        grouped = group_rows_by_investor(rows)
        for _, group_rows in grouped.items():
            first = group_rows[0]
            investor_subject = format_investor_subject(first.get("adsoyad", ""), first.get("kimlik_no", ""))
            paragraphs.append(
                f"İlgi'de kayıtlı yazınızda yer alan yatırımcının TCKN bilgisi ile Merkezimiz kayıtları üzerinde yapılan incelemede, "
                f"anılan şahsın Merkezimiz nezdinde herhangi bir emanet, hak ve alacağına rastlanılmamıştır (Ek)."
            )
        return paragraphs

    identity = extract_report_identity(pdf_path)
    investor_subject = format_investor_subject(identity.get("adsoyad", ""), identity.get("kimlik_no", ""))
    paragraphs.append(
        f"SPK'nun 83’üncü maddesinin dördüncü fıkrası kapsamında "
        f"{investor_subject} adlı yatırımcının Merkezimiz uhdesinde herhangi bir hak, alacak veya emanet kaydına rastlanılmamıştır."
    )
    return paragraphs


def create_docx(paragraphs: List[str], output_path: Path) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for text in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.save(output_path)


# ============================================================
# Debug dosyası
# ============================================================

def save_debug_file(pdf_path: Path, debug_lines: List[str], rows: List[Dict[str, str]]) -> Path:
    debug_path = pdf_path.with_name(f"{pdf_path.stem}_debug_satirlar.txt")

    with open(debug_path, "w", encoding="utf-8") as f:
        f.write("=== PDF'TEN SATIR BAZINDA OKUNAN İÇERİK ===\n\n")

        for i, line in enumerate(debug_lines, start=1):
            f.write(f"{i:03d}. {line}\n")

        f.write("\n\n=== YAKALANAN BAKİYE SATIRLARI ===\n\n")

        if not rows:
            f.write("Bakiye satırı yakalanamadı.\n")
        else:
            for i, row in enumerate(rows, start=1):
                f.write(f"--- Kayıt {i} ---\n")
                for key, value in row.items():
                    f.write(f"{key}: {value}\n")
                f.write("\n")

    return debug_path


# ============================================================
# Ana arayüz
# ============================================================

def main():
    root = tk.Tk()
    root.withdraw()

    user_answer = messagebox.askyesnocancel(
        "Bakiye Beyanı",
        "Kullanıcı beyanına göre bu müşterinin bakiyesi var mı?"
    )

    if user_answer is None:
        messagebox.showinfo("İptal", "İşlem iptal edildi.")
        return

    pdf_file = filedialog.askopenfilename(
        title="Bakiye raporu PDF dosyasını seçiniz",
        filetypes=[("PDF dosyaları", "*.pdf")]
    )

    if not pdf_file:
        messagebox.showinfo("İptal", "PDF seçilmedi.")
        return

    pdf_path = Path(pdf_file)

    try:
        rows, debug_lines = extract_balance_rows(pdf_path)
        positive_rows = [row for row in rows if has_positive_balance(row)]

        debug_file = save_debug_file(pdf_path, debug_lines, rows)

        pdf_has_positive_balance = len(positive_rows) > 0

        if not rows:
            messagebox.showinfo(
                "Bakiye Raporu Boş",
                "PDF içinde bakiye satırı bulunamadı.\n\n"
                "Boş rapor metni oluşturulacak.\n\n"
                f"Debug dosyası:\n{debug_file}"
            )

        if not user_answer and pdf_has_positive_balance:
            messagebox.showwarning(
                "Uyarı",
                "Kullanıcı bakiyesi yok dedi; ancak PDF içinde pozitif bakiye bulundu.\n\n"
                "Word yazısı PDF verisine göre oluşturulacak."
            )

        if user_answer and not pdf_has_positive_balance:
            messagebox.showwarning(
                "Uyarı",
                "Kullanıcı bakiyesi var dedi; ancak PDF içinde pozitif bakiye bulunamadı.\n\n"
                "Boş rapor metni oluşturulacak.\n\n"
                f"Debug dosyası:\n{debug_file}"
            )

        if pdf_has_positive_balance:
            paragraphs = build_word_paragraphs(positive_rows)
        else:
            messagebox.showinfo(
                "Bakiye Yok",
                "PDF içinde pozitif bakiye bulunamadı.\n\n"
                "Boş rapor metni oluşturulacak.\n\n"
                f"Debug dosyası:\n{debug_file}"
            )
            paragraphs = build_empty_report_paragraphs(rows, pdf_path)

        default_name = (
            f"{pdf_path.stem}_bakiye_yazisi_"
            f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )

        output_file = filedialog.asksaveasfilename(
            title="Word dosyasını kaydet",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word dosyası", "*.docx")]
        )

        if not output_file:
            messagebox.showinfo("İptal", "Word dosyası kaydedilmedi.")
            return

        output_path = Path(output_file)
        create_docx(paragraphs, output_path)

        preview = "\n\n".join(paragraphs)

        messagebox.showinfo(
            "İşlem Tamamlandı",
            f"Word dosyası oluşturuldu:\n\n{output_path}\n\n"
            f"Debug dosyası:\n{debug_file}\n\n"
            f"Oluşturulan metin:\n\n{preview}"
        )

        try:
            if os.name == "nt":
                os.startfile(output_path)
        except Exception:
            pass

    except Exception as exc:
        messagebox.showerror(
            "Hata",
            f"İşlem sırasında hata oluştu:\n\n{exc}"
        )
        raise


if __name__ == "__main__":
    main()
